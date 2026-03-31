from fastapi import APIRouter, HTTPException, Header
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Dict, Optional
import google.generativeai as genai
import os
import json
import re
import io
import httpx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

router = APIRouter()

# Configuración API
api_key = os.getenv("GOOGLE_API_KEY")
if api_key:
    genai.configure(api_key=api_key)

API_BASE_URL = os.getenv("API_BASE_URL", "http://localhost:8000")
CONTEXTO_FALLBACK = "UBICACIÓN: Chile."

# --- MODELOS DE DATOS ---
class NeeRequest(BaseModel):
    grade: str
    subject: str
    diagnosis: str      # Ej: TEA, TDAH, FIL
    barrier: str        # Ej: Le molestan los ruidos, no lee fluido...
    activity: str       # La actividad original

class Estrategias(BaseModel):
    acceso: str         # Estrategias previas / ambiente
    actividad: str      # Adaptación de la tarea
    evaluacion: str     # Forma de evaluar

class DownloadRequest(BaseModel):
    grade: str
    subject: str
    diagnosis: str
    barrier: str
    activity: str
    estrategias: Estrategias
    estrategias: Estrategias
    dua_principles: str # Breve explicación de qué principios DUA se usan

# --- MODELOS PARA GENERACIÓN DUA ---
class DUARequest(BaseModel):
    planificacion_original: str
    contexto_estudiante: str

class VariantesDUA(BaseModel):
    visual_espacial: str
    kinestesica: str
    focalizada: str


# --- PROMPT ESPECIALISTA EN INCLUSIÓN ---
SYSTEM_PROMPT = """
ROL: Eres un Experto en Educación Diferencial y DUA (Diseño Universal de Aprendizaje).
MISIÓN: Recibes una actividad de clase y un perfil de estudiante (NEE Transitoria o Permanente). Debes generar adecuaciones curriculares prácticas para el aula.

ENTRADA:
- Curso y Asignatura
- Diagnóstico (Ej: TEA, TDAH)
- Barrera Específica (Lo que le cuesta al estudiante)
- Actividad Original

TU TAREA (JSON Estricto):
Genera una respuesta estructurada con enfoque de AULA (Práctico, no burocrático):

1. Principios DUA: Breve línea citando qué principio aplicas (Representación, Acción/Expresión, Compromiso).
2. Estrategia de Acceso (Antes): Qué cambiar en el ambiente, materiales o ubicación.
3. Adecuación de la Actividad (Durante): Modificación de la instrucción o tarea.
4. Evaluación Diversificada (Cierre): Cómo puede demostrar aprendizaje de otra forma.

FORMATO JSON:
{
  "dua_principles": "Texto breve...",
  "estrategias": {
    "acceso": "Texto detallado...",
    "actividad": "Texto detallado...",
    "evaluacion": "Texto detallado..."
  }
}
"""

@router.post("/nee/generate")
async def generate_nee_strategies(req: NeeRequest, authorization: Optional[str] = Header(None)):
    try:
        # Contexto institucional dinámico
        contexto_institucional = CONTEXTO_FALLBACK
        if authorization:
            try:
                async with httpx.AsyncClient(timeout=5.0) as client:
                    ctx_resp = await client.get(
                        f"{API_BASE_URL}/profile/context",
                        headers={"Authorization": authorization}
                    )
                    if ctx_resp.status_code == 200:
                        block = ctx_resp.json().get("context_block", "")
                        if block:
                            contexto_institucional = block
            except Exception:
                pass

        model = genai.GenerativeModel('gemini-2.5-flash')
        
        prompt = f"""
        CONTEXTO INSTITUCIONAL: {contexto_institucional}
        
        CONTEXTO DEL CASO:
        - Curso: {req.grade}
        - Asignatura: {req.subject}
        - Diagnóstico: {req.diagnosis}
        - Barrera Principal: "{req.barrier}"
        - Actividad Original: "{req.activity}"

        Genera las adecuaciones prácticas en JSON. Sé empático y concreto.
        """
        
        response = model.generate_content(SYSTEM_PROMPT + prompt)
        
        # Limpieza de JSON
        texto_limpio = response.text
        if "```" in texto_limpio:
            texto_limpio = re.sub(r"```json\s*", "", texto_limpio)
            texto_limpio = re.sub(r"```\s*$", "", texto_limpio)
        
        data = json.loads(texto_limpio.strip())
        return data

    except Exception as e:
        print(f"❌ Error NEE: {e}")
        return {
            "dua_principles": "Error generando análisis.",
            "estrategias": {
                "acceso": "Intenta simplificar la descripción de la barrera.",
                "actividad": "",
                "evaluacion": ""
            }
            }


# --- GENERADOR DUA (NUEVO endpoint) ---
@router.post("/nee/generate-dua")
async def generate_dua_variants(req: DUARequest):
    try:
        print(f"🧩 [NEE] Generando DUA para: {req.contexto_estudiante}")
        
        prompt = f"""
        ACTÚA COMO ESPECIALISTA EN EDUCACIÓN DIFERENCIAL.
        
        INPUT:
        1. PLANIFICACIÓN BASE:
        "{req.planificacion_original}"
        
        2. CONTEXTO DEL ESTUDIANTE:
        "{req.contexto_estudiante}"
        
        TAREA:
        Genera 3 variantes de adecuación curricular para la actividad principal de esta planificación.
        
        SALIDA JSON ESTRICTA:
        {{
            "visual_espacial": "Descripción de estrategia con apoyo visual/gráfico...",
            "kinestesica": "Descripción de estrategia manipulativa/corporal...",
            "focalizada": "Estrategia específica y simplificada para el perfil descrito..."
        }}
        """

        model = genai.GenerativeModel('gemini-2.5-flash', generation_config={"response_mime_type": "application/json"})
        response = model.generate_content(prompt)
        
        return json.loads(response.text)

    except Exception as e:
        print(f"❌ Error DUA: {e}")
        return {
            "visual_espacial": "Error generando variante visual.",
            "kinestesica": "Error generando variante kinestésica.",
            "focalizada": "Error generando variante focalizada."
        }


# --- EXPORTACIÓN A WORD ---
@router.post("/nee/download")
async def download_nee_docx(data: DownloadRequest):
    try:
        doc = Document()
        
        # Logo y Título
        t = doc.add_table(rows=1, cols=2)
        t.autofit = False
        t.columns[0].width = Inches(1.2)
        t.columns[1].width = Inches(5.3)
        
        # Logo
        cell_logo = t.cell(0, 0)
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        logo_path = os.path.join(base, "assets", "logo_profeic.svg.png")
        p_logo = cell_logo.paragraphs[0]
        run_logo = p_logo.add_run()
        if os.path.exists(logo_path):
            try:
                run_logo.add_picture(logo_path, width=Inches(1.0))
            except Exception:
                run_logo.add_text("ProfeIC")
        else:
            run_logo.add_text("ProfeIC")

        # Título
        cell_info = t.cell(0, 1)
        p_info = cell_info.paragraphs[0]
        p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p_info.add_run("PROFE IC\n")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(43, 84, 110)
        p_info.add_run("Asistente de Inclusión & DUA\n").font.size = Pt(10)
        doc.add_paragraph()

        # Datos del Estudiante
        doc.add_heading('1. Contexto del Estudiante', level=1)
        doc.add_paragraph(f"Curso: {data.grade} | Asignatura: {data.subject}")
        p = doc.add_paragraph()
        p.add_run("Diagnóstico/Condición: ").bold = True
        p.add_run(data.diagnosis)
        p = doc.add_paragraph()
        p.add_run("Barrera Detectada: ").bold = True
        p.add_run(data.barrier)

        # Actividad Original
        doc.add_heading('2. Actividad Original', level=1)
        doc.add_paragraph(data.activity)

        # Adecuaciones
        doc.add_heading('3. Estrategias de Adecuación Curricular', level=1)
        
        p = doc.add_paragraph()
        p.add_run("Enfoque DUA: ").bold = True
        p.add_run(data.dua_principles).italic = True

        doc.add_heading('A. Adecuación de Acceso (Preparación)', level=2)
        doc.add_paragraph(data.estrategias.acceso)

        doc.add_heading('B. Adecuación de la Actividad (Desarrollo)', level=2)
        doc.add_paragraph(data.estrategias.actividad)

        doc.add_heading('C. Evaluación Diversificada', level=2)
        doc.add_paragraph(data.estrategias.evaluacion)

        # Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "Documento generado por PROFE IC - Recurso creado con Inteligencia Aumentada de la Plataforma ProfeIC"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Adecuacion_Curricular.docx"}
        )

    except Exception as e:
        print(f"❌ Error DOCX NEE: {e}")
        raise HTTPException(status_code=500, detail="Error generando documento")