import os
from typing import Literal, Optional
from fastapi import APIRouter, HTTPException, File, UploadFile, Depends
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel
import google.generativeai as genai
from supabase import create_client
from routers.deps import get_current_user_id

router = APIRouter(prefix="/api/v1/pme", tags=["PME"])

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY") or os.getenv("SUPABASE_KEY", "")

try:
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception:
    supabase = None

genai.configure(api_key=os.getenv("GOOGLE_API_KEY", ""))


def get_model():
    return genai.GenerativeModel(
        "gemini-2.5-flash",
        generation_config={"response_mime_type": "application/json"}
    )


@router.get("/analisis-progreso")
async def analisis_progreso(user_id: str = Depends(get_current_user_id)):
    """Analiza las metas y acciones del PME del colegio del usuario usando IA."""
    try:
        # 1. Validar e identificar la escuela del usuario
        profile_res = supabase.table("profiles").select("school_id").eq("id", user_id).single().execute()
        school_id = profile_res.data.get("school_id") if profile_res.data else None
        if not school_id:
            return JSONResponse(status_code=403, content={"alerta": "El usuario no tiene una institución asignada."})

        # 2. Filtrar consultas estrictamente por school_id
        pme_res = supabase.table("pme_actions").select("*").eq("school_id", school_id).execute()
        actions = pme_res.data or []
        
        goals_res = supabase.table("strategic_goals").select("*").eq("school_id", school_id).not_.is_("pme_action_link", "null").execute()
        goals = goals_res.data or []

        if not goals:
            return JSONResponse({"alerta": "No hay metas operativas vinculadas actualmente. El PME oficial necesita despliegue operativo."})

        # Preparamos un prompt para gemini
        prompt = f"""
        Actúa como un auditor experto en gestión escolar y analista de datos educativos. 
        Te entregaré las acciones oficiales del PME y las metas operativas vinculadas.
        Haz un breve análisis ejecutivo identificando cuellos de botella ("estancamientos") o dando retroalimentación.
        Devuelve SOLO un JSON con la estructura {{"alerta": "tu texto aquí"}}. Sin markdown, solo JSON puro.
        
        DATOS:
        Acciones PME: {actions}
        Metas Operativas Vinculadas: {goals}
        """

        model = get_model()
        response = await model.generate_content_async(prompt)
        text = response.text.replace("```json", "").replace("```", "").strip()
        import json
        try:
            parsed = json.loads(text)
            return JSONResponse(parsed)
        except json.JSONDecodeError:
            return JSONResponse({"alerta": text})

    except Exception as e:
        print(f"[ERROR] PME Analysis: {e}")
        return JSONResponse({"alerta": "Error temporal al analizar el ritmo de progreso. Consulte más tarde."})


@router.get("/exportar")
async def exportar_evidencia(user_id: str = Depends(get_current_user_id)):
    """Exporta el reporte de cumplimiento filtrado por la institución del usuario."""
    try:
        # 1. Validar e identificar la escuela del usuario
        profile_res = supabase.table("profiles").select("school_id").eq("id", user_id).single().execute()
        school_id = profile_res.data.get("school_id") if profile_res.data else None
        if not school_id:
            raise HTTPException(status_code=403, detail="El usuario no tiene una institución asignada.")

        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        doc = Document()
        
        # Titulo
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("EVIDENCIA DE CUMPLIMIENTO PME\n(Acciones Oficiales y Metas Operativas)")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph() # Spacer

        # 2. Consultar datos exclusivamente para school_id
        pme_res = supabase.table("pme_actions").select("*").eq("school_id", school_id).execute()
        actions = pme_res.data or []
        
        goals_res = supabase.table("strategic_goals").select("*").eq("school_id", school_id).not_.is_("pme_action_link", "null").execute()
        goals = goals_res.data or []
        
        phases_res = supabase.table("implementation_phases").select("*").execute()  # las fases se asocian a metas
        phases = phases_res.data or []
        
        inds_res = supabase.table("indicators").select("*").execute()
        inds = inds_res.data or []

        if not actions:
            doc.add_paragraph("No hay acciones oficiales registradas en el PME.")
        
        for action in actions:
            p_action = doc.add_paragraph()
            r_action = p_action.add_run(f"Acción PME: {action.get('title', 'Sin Título')} ({action.get('dimension', 'Sin Dimensión')})")
            r_action.bold = True
            r_action.font.size = Pt(12)

            linked_goals = [g for g in goals if g.get("pme_action_link") == action.get("id")]
            if not linked_goals:
                p_none = doc.add_paragraph("  - Sin metas operativas vinculadas.")
                p_none.style = 'List Bullet'
            else:
                for goal in linked_goals:
                    p_goal = doc.add_paragraph(f"Meta: {goal.get('title')}")
                    p_goal.style = 'List Bullet'
                    
                    goal_phases = [p for p in phases if p.get("goal_id") == goal.get("id")]
                    for phase in goal_phases:
                        doc.add_paragraph(f"Fase: {phase.get('title')}", style='List Bullet 2')
                        
                        phase_inds = [i for i in inds if i.get("phase_id") == phase.get("id")]
                        for ind in phase_inds:
                            doc.add_paragraph(f"{ind.get('description')} - Progreso: {ind.get('current_value')}/{ind.get('target_value')}", style='List Bullet 3')
            
            doc.add_paragraph() # Spacer

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        headers = {
            "Content-Disposition": 'attachment; filename="Evidencia_Cumplimiento_PME.docx"'
        }
        return Response(
            content=file_stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] PME Export: {e}")
        return JSONResponse(status_code=500, content={"alerta": "Error al exportar documento."})


@router.post("/importar-documento")
async def importar_documento_pme(file: UploadFile = File(...), user_id: str = Depends(get_current_user_id)):
    """Importa e interpreta un archivo de reporte PME oficial."""
    try:
        import io
        from pypdf import PdfReader

        # Validar tipo de archivo
        if not file.filename.lower().endswith(".pdf"):
            return JSONResponse(status_code=400, content={"message": "Formato inválido. Por favor sube un archivo PDF."})
        
        # Leer archivo
        content = await file.read()
        pdf = PdfReader(io.BytesIO(content))
        text_content = ""
        for page in pdf.pages:
            text_content += page.extract_text() + "\n"
            
        if not text_content.strip():
            return JSONResponse(status_code=400, content={"message": "No se pudo extraer texto del PDF (probablemente es escaneado)."})

        # Reducir texto si es demasiado gigante
        if len(text_content) > 100000:
            text_content = text_content[:100000]

        prompt = f"""Eres un analista de datos escolares. A continuación recibirás el texto extraído de un reporte oficial del Ministerio de Educación (PME). 
        Tu tarea es extraer y clasificar la información en la siguiente estructura JSON exacta, omitiendo textos de relleno:
        
        {{
          "identidad": {{
            "vision": "Texto de la visión o 'No especificada'",
            "mision": "Texto de la misión o 'No especificada'",
            "sellos": ["Sello 1", "Sello 2"]
          }},
          "recursos": {{
            "total": numero_entero,
            "sep": numero_entero,
            "pie": numero_entero
          }},
          "estrategias": [
            {{
              "dimension": "Ej: Gestión Pedagógica",
              "objetivo": "Texto del objetivo",
              "estrategia": "Texto de la estrategia"
            }}
          ],
          "diagnostico": {{
            "fortalezas": ["Fortaleza 1", "Fortaleza 2"],
            "oportunidades_mejora": ["Mejora 1", "Mejora 2"]
          }}
        }}
        
        Asegúrate de devolver SOLO este JSON válido.
        
        === TEXTO DEL REPORTE ===
        {text_content}
        """

        # Enviar a Gemini con JSON mode
        import json
        model = get_model()
        response = await model.generate_content_async(prompt)
        
        result_text = response.text.replace("```json", "").replace("```", "").strip()
        parsed_data = json.loads(result_text)
        
        return JSONResponse(content=parsed_data)

    except Exception as e:
        print(f"[ERROR] Importar PME: {e}")
        return JSONResponse(status_code=500, content={"message": "Ocurrió un error al procesar el archivo. Intente nuevamente."})


class ConsolidateRequest(BaseModel):
    identidad: dict = None
    recursos: dict = None
    estrategias: list = []
    diagnostico: dict = None
    school_id: str = None  # Se validará contra el token del usuario


@router.post("/consolidar")
async def consolidar_pme(req: ConsolidateRequest, user_id: str = Depends(get_current_user_id)):
    """Guarda la información PME extraída garantizando el aislamiento multitenant."""
    try:
        # 1. Validar e identificar la escuela del usuario
        profile_res = supabase.table("profiles").select("school_id").eq("id", user_id).single().execute()
        school_id = profile_res.data.get("school_id") if profile_res.data else None
        if not school_id:
            raise HTTPException(status_code=403, detail="El usuario no tiene una institución asignada.")
            
        # 2. Prevenir Spoofing: validar o forzar school_id de la sesión
        target_school_id = school_id
        if req.school_id and str(req.school_id) != str(school_id):
            raise HTTPException(status_code=403, detail="Violación de multitenancy: El school_id solicitado no corresponde.")
        
        # 3. Guardar Identidad e Información Financiera
        if req.identidad or req.recursos:
            info_to_save = {
                "mission": req.identidad.get("mision") if req.identidad else None,
                "vision": req.identidad.get("vision") if req.identidad else None,
                "sellos": req.identidad.get("sellos", []) if req.identidad else [],
                "budget_sep": req.recursos.get("sep", 0) if req.recursos else 0,
                "budget_pie": req.recursos.get("pie", 0) if req.recursos else 0,
                "budget_total": req.recursos.get("total", 0) if req.recursos else 0,
                "school_id": target_school_id,
                "updated_at": "now()"
            }
            # Upsert
            year = 2026
            supabase.table("pme_institutional_info").upsert(
                {**info_to_save, "academic_year": year},
                on_conflict="academic_year"
            ).execute()

        # 4. Guardar Estrategias (Acciones Oficiales)
        if req.estrategias:
            actions_to_insert = []
            for estr in req.estrategias:
                actions_to_insert.append({
                    "school_id": target_school_id,
                    "dimension": estr.get("dimension"),
                    "title": estr.get("objetivo"),
                    "description": estr.get("estrategia"),
                    "academic_year": 2026,
                    "status": "active"
                })
            
            if actions_to_insert:
                supabase.table("pme_actions").insert(actions_to_insert).execute()

        return JSONResponse(content={"message": "Plan consolidado con éxito en la base de datos oficial."})

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Consolidar PME: {e}")
        return JSONResponse(status_code=500, content={"message": f"Fallo al guardar: {str(e)}"})


class ValidarSmartRequest(BaseModel):
    tipo: Literal["meta", "indicador"]
    texto: str
    contexto: Optional[str] = None


@router.post("/validar-smart")
async def validar_smart(req: ValidarSmartRequest, user_id: str = Depends(get_current_user_id)):
    """Audita un objetivo o indicador PME usando los criterios SMART y sugiere una propuesta optimizada."""
    if not req.texto.strip():
        raise HTTPException(status_code=400, detail="El texto a evaluar no puede estar vacío.")
    
    try:
        prompt = f"""
        Actúa como un auditor experto en planificación estratégica escolar y gestión pedagógica (PME MINEDUC).
        Tu tarea es analizar el siguiente {req.tipo} utilizando estrictamente los criterios SMART:
        - **S (Specific - Específico)**: ¿El objetivo/indicador está delimitado con precisión y es claro?
        - **M (Measurable - Medible)**: ¿Contiene un criterio cuantitativo, porcentaje o meta numérica para medir su éxito?
        - **A (Achievable - Alcanzable)**: ¿Es realista considerando las capacidades usuales de un establecimiento escolar?
        - **R (Relevant - Relevante)**: ¿Aporta valor pedagógico y se alinea con la mejora de aprendizajes?
        - **T (Time-bound - Acotado en el Tiempo)**: ¿Define un plazo temporal explícito para su cumplimiento?

        Texto a evaluar ({req.tipo}): "{req.texto}"
        {f'Contexto del plan: "{req.contexto}"' if req.contexto else ''}

        Debes evaluar cada uno de los 5 criterios anteriores. Para cada criterio asigna uno de los siguientes estados: "ok", "mejorable", "insuficiente".
        Calcula un puntaje global (score) entre 0 y 100 basado en el cumplimiento.
        Proporciona un listado de sugerencias concretas de mejora en español.
        Por último, ofrece una redacción SMART alternativa y optimizada que sirva de propuesta formal (rewritten_proposal), la cual debe incluir un indicador medible y un plazo de tiempo plausible de forma explícita.

        Retorna EXCLUSIVAMENTE un objeto JSON válido con la siguiente estructura exacta:
        {{
          "is_smart": {"true" if "true" else "false"},
          "score": 85,
          "evaluation": {{
            "specific": {{
              "status": "ok",
              "details": "Detalle del criterio..."
            }},
            "measurable": {{
              "status": "ok",
              "details": "Detalle del criterio..."
            }},
            "achievable": {{
              "status": "ok",
              "details": "Detalle del criterio..."
            }},
            "relevant": {{
              "status": "ok",
              "details": "Detalle del criterio..."
            }},
            "time_bound": {{
              "status": "ok",
              "details": "Detalle del criterio..."
            }}
          }},
          "suggestions": [
            "Sugerencia 1",
            "Sugerencia 2"
          ],
          "rewritten_proposal": "Redacción SMART..."
        }}
        """

        model = get_model()
        response = await model.generate_content_async(prompt)
        text = response.text.replace("```json", "").replace("```", "").strip()
        
        import json
        try:
            parsed = json.loads(text)
            return JSONResponse(parsed)
        except json.JSONDecodeError:
            # Fallback simple
            return JSONResponse({
                "is_smart": False,
                "score": 50,
                "evaluation": {
                    "specific": {"status": "mejorable", "details": "Error al parsear el análisis detallado."},
                    "measurable": {"status": "mejorable", "details": "Error al parsear el análisis detallado."},
                    "achievable": {"status": "mejorable", "details": "Error al parsear el análisis detallado."},
                    "relevant": {"status": "mejorable", "details": "Error al parsear el análisis detallado."},
                    "time_bound": {"status": "mejorable", "details": "Error al parsear el análisis detallado."}
                },
                "suggestions": ["Por favor vuelve a intentar para obtener sugerencias detalladas."],
                "rewritten_proposal": req.texto
            })
    except Exception as e:
        print(f"[ERROR] Validar SMART: {e}")
        return JSONResponse(status_code=500, content={"message": f"Error interno del servidor: {str(e)}"})

