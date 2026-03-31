from fastapi import APIRouter, HTTPException, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional, Union
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import os
import re
import json
import json
import traceback
import base64
import uuid
import unicodedata
from urllib.parse import quote

download_cache = {}

router = APIRouter()

def es_evaluacion(content: Any) -> bool:
    if isinstance(content, list) and len(content) > 0 and isinstance(content[0], dict) and ("stem" in content[0] or "type" in content[0]):
        return True
    if isinstance(content, dict):
        if "items" in content and isinstance(content.get("items"), list): return True
        if "student_version" in content and isinstance(content.get("student_version"), dict) and "items" in content["student_version"]: return True
        if "teacher_guide" in content and "answers" in content.get("teacher_guide", {}): return True
        if "quantities" in content and "points" in content: return True
    return False

def get_safe_headers(filename: str) -> dict:
    # 1. Normalizar ASCII para navegadores viejos/estrictos
    safe_ascii = unicodedata.normalize('NFKD', filename).encode('ASCII', 'ignore').decode('utf-8')
    safe_ascii = re.sub(r'[^\w\s\.-]', '', safe_ascii).replace(' ', '_')
    # 2. Codificar UTF-8 RFC 5987 para navegadores modernos
    safe_utf8 = quote(filename.replace(' ', '_'))
    return {
        "Content-Disposition": f"attachment; filename=\"{safe_ascii}\"; filename*=UTF-8''{safe_utf8}"
    }

# ==========================================
# 1. MODELOS DE DATOS
# ==========================================

class ExportRequest(BaseModel):
    titulo: str
    descripcion: str
    nivel: str
    asignatura: str
    oa: str
    actividad: str
    puntaje_total: int
    tabla: List[Dict[str, Any]]

class PlanExportRequest(BaseModel):
    titulo_unidad: str
    estrategia: str
    nivel: str
    asignatura: str
    oas: List[str]
    clases: List[Dict[str, Any]]

class AssessmentItem(BaseModel):
    type: str
    dok_level: int
    points: int
    stem: str
    options: Optional[List[Union[str, Dict[str, Any]]]] = None
    correct_answer: Optional[bool] = None
    rubric_hint: Optional[str] = None

class AssessmentExportRequest(BaseModel):
    title: str
    description: str
    grade: str
    subject: str
    items: List[AssessmentItem]
    teacher_guide: Optional[Dict[str, Any]] = None # <--- Nueva clave importada

class ElevatorExportRequest(BaseModel):
    grade: str
    subject: str
    activity: str
    dok_actual: str
    diagnostico: str
    escalera: List[Dict[str, Any]]
    propuestas: Dict[str, str]

class LecturaPregunta(BaseModel):
    id: str
    nivel_taxonomico: str
    pregunta: str
    alternativas: List[str]
    respuesta_correcta: str
    justificacion: str

class LecturaInteligenteExportRequest(BaseModel):
    nivel: str
    asignatura: str
    oa: str
    texto: str
    preguntas: List[LecturaPregunta]
    tipo_documento: str # "estudiante" o "profesor"

class GenericExportRequest(BaseModel):
    titulo_unidad: str
    nivel: str
    asignatura: str
    contenido: Union[Dict[str, Any], List[Any]]

class PaqueteExportRequest(BaseModel):
    documentos: List[GenericExportRequest]

# ==========================================
# 2. UTILIDADES DE DISEÑO
# ==========================================

def limpiar_latex_para_word(texto: str) -> str:
    """Convierte LaTeX a Unicode legible en Word. NO elimina los símbolos, los transforma."""
    if not texto: return ""
    texto = str(texto)
    
    # 1. Convertir raíces cuadradas: \sqrt{x} -> √x | \sqrt{xy} -> √(xy)
    texto = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', texto)
    texto = re.sub(r'\\sqrt\s+(\S+)', r'√\1', texto)
    
    # 2. Convertir fracciones: \frac{a}{b} -> a/b
    texto = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', texto)
    
    # 3. Potencias: x^{2} -> x², x^2 -> x²
    superscript_map = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹'}
    def replace_superscript(m):
        exp = m.group(1)
        return ''.join(superscript_map.get(c, c) for c in exp)
    texto = re.sub(r'\^\{([^}]*)\}', replace_superscript, texto)
    texto = re.sub(r'\^([0-9])', lambda m: superscript_map.get(m.group(1), m.group(1)), texto)
    
    # 4. Subíndices: x_{2} -> x₂
    subscript_map = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉'}
    def replace_subscript(m):
        sub = m.group(1)
        return ''.join(subscript_map.get(c, c) for c in sub)
    texto = re.sub(r'_\{([^}]*)\}', replace_subscript, texto)
    
    # 5. Letras griegas y símbolos matemáticos comunes
    replacements = {
        r'\\pi': 'π', r'\\Pi': 'Π',
        r'\\alpha': 'α', r'\\beta': 'β', r'\\gamma': 'γ', r'\\delta': 'δ',
        r'\\theta': 'θ', r'\\lambda': 'λ', r'\\mu': 'μ', r'\\sigma': 'σ',
        r'\\Sigma': 'Σ', r'\\omega': 'ω', r'\\Omega': 'Ω',
        r'\\infty': '∞', r'\\approx': '≈', r'\\neq': '≠',
        r'\\leq': '≤', r'\\geq': '≥', r'\\pm': '±',
        r'\\times': '×', r'\\div': '÷', r'\\cdot': '·',
        r'\\in': '∈', r'\\notin': '∉', r'\\subset': '⊂', r'\\cup': '∪', r'\\cap': '∩',
        r'\\forall': '∀', r'\\exists': '∃', r'\\neg': '¬',
        r'\\rightarrow': '→', r'\\leftarrow': '←', r'\\Rightarrow': '⇒',
        r'\\angle': '∠', r'\\triangle': '△', r'\\degree': '°',
        r'\\%': '%',
        r'\\sin': 'sin', r'\\cos': 'cos', r'\\tan': 'tan',
        r'\\log': 'log', r'\\ln': 'ln', r'\\exp': 'exp',
    }
    for latex, unicode_char in replacements.items():
        texto = re.sub(latex, unicode_char, texto)
    
    # 6. Quitar delimitadores de fórmulas $...$ y $$...$$ (el contenido ya fue transformado)
    texto = re.sub(r'\$\$([^$]*)\$\$', r'\1', texto)
    texto = re.sub(r'\$([^$]*)\$', r'\1', texto)
    
    # 7. Limpiar backslashes residuales (mantener el texto del comando LaTeX no reconocido)
    texto = re.sub(r'\\([a-zA-Z]+)', r'\1 ', texto)
    texto = texto.replace('\\', '')
    
    # 8. Limpiar llaves sobrantes
    texto = texto.replace('{', '').replace('}', '')
    
    return texto.strip()

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def style_header_cell(cell, text, bg_color="2B546E", text_color="FFFFFF"):
    cell.text = text
    set_cell_background(cell, bg_color)
    if cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        try:
            run.font.color.rgb = RGBColor.from_string(text_color)
        except Exception:
            pass
        run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_header_logo(doc, asignatura, nivel, titulo_extra=""):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(5.3)
    
    cell_logo = table.cell(0, 0)
    p = cell_logo.paragraphs[0]
    run = p.add_run()
    
    from pathlib import Path
    base_dir = Path(__file__).resolve().parent.parent
    logo_path = base_dir / 'assets' / 'logo_profeic.svg.png'
    
    if logo_path.exists():
        try:
            run.add_picture(str(logo_path), width=Inches(1.0))
        except Exception:
            run.add_text("ProfeIC")
    else:
        run.add_text("ProfeIC")

    cell_info = table.cell(0, 1)
    p_info = cell_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_info.add_run("PROFE IC\n")
    r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = RGBColor(43, 84, 110)
    
    r2 = p_info.add_run("Recurso creado con Inteligencia Aumentada de la Plataforma ProfeIC\n")
    r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(100, 100, 100)
    
    p_info.add_run(f"Asignatura: {asignatura} | Nivel: {nivel}\n").font.size = Pt(10)
    if titulo_extra:
        p_info.add_run(f"{titulo_extra}").italic = True

    doc.add_paragraph() 

# ==========================================
# 3. MOTORES DE RENDERIZADO (Lógica Visual)
# ==========================================

def renderizar_planificacion(doc, data):
    """Motor Visual para Planificaciones"""
    titulo = data.get('titulo_unidad_creativo') or data.get('titulo_unidad') or "Planificación"
    h = doc.add_heading(limpiar_latex_para_word(titulo), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    est = data.get('estrategia_aprendizaje_sentencia') or data.get('estrategia')
    if est:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Estrategia: {limpiar_latex_para_word(est)}")
        r.italic = True; r.font.color.rgb = RGBColor(43, 84, 110)

    # Clases
    clases = data.get('planificacion_clases') or data.get('clases') or []
    for clase in clases:
        doc.add_paragraph()
        # Header Clase
        t = doc.add_table(rows=1, cols=1); t.autofit = False; t.columns[0].width = Inches(6.5)
        c = t.cell(0,0); set_cell_background(c, "2B546E")
        p = c.paragraphs[0]; r = p.add_run(f"CLASE {clase.get('numero_clase','?')}"); r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(12)
        if clase.get('foco_pedagogico'): p.add_run(f" | Foco: {limpiar_latex_para_word(clase.get('foco_pedagogico'))}").font.color.rgb=RGBColor(255,255,255)
        
        # Momentos
        tm = doc.add_table(rows=2, cols=3); tm.style = 'Table Grid'
        style_header_cell(tm.cell(0,0), "INICIO", "F2AE60"); style_header_cell(tm.cell(0,1), "DESARROLLO", "F2AE60"); style_header_cell(tm.cell(0,2), "CIERRE", "F2AE60")
        cont = clase.get('contenido_editable', {})
        tm.cell(1,0).text = limpiar_latex_para_word(cont.get('inicio',''))
        tm.cell(1,1).text = limpiar_latex_para_word(cont.get('desarrollo',''))
        tm.cell(1,2).text = limpiar_latex_para_word(cont.get('cierre',''))

        # Recursos
        doc.add_paragraph()
        tr = doc.add_table(rows=1, cols=1); tr.style = 'Table Grid'
        tr.cell(0,0).text = f"Recursos: {limpiar_latex_para_word(clase.get('recurso_practica', ''))}\nTicket: {limpiar_latex_para_word(clase.get('ticket_salida', ''))}"

def renderizar_rubrica(doc, data):
    """Motor Visual para Rúbricas"""
    # Título y Descripción
    titulo = data.get('titulo') or "Rúbrica de Evaluación"
    h = doc.add_heading(limpiar_latex_para_word(titulo), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if data.get('descripcion'):
        p = doc.add_paragraph(limpiar_latex_para_word(data.get('descripcion')))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.italic = True
    
    doc.add_paragraph()
    
    # Metadata (OA, Actividad)
    if data.get('oaDescripcion') or data.get('oa'):
        p = doc.add_paragraph()
        p.add_run("Objetivo de Aprendizaje: ").bold = True
        p.add_run(limpiar_latex_para_word(data.get('oaDescripcion') or data.get('oa')))
    
    if data.get('actividad'):
        p = doc.add_paragraph()
        p.add_run("Actividad / Producto: ").bold = True
        p.add_run(limpiar_latex_para_word(data.get('actividad')))
        
    doc.add_paragraph()

    # TABLA RÚBRICA
    tabla_datos = data.get('tabla', [])
    if tabla_datos:
        # Definir columnas: Criterio + 4 Niveles
        t = doc.add_table(rows=1, cols=5)
        t.style = 'Table Grid'
        t.autofit = False
        
        # Headers con Colores
        headers = ["Criterio & Peso", "Insuficiente", "Elemental", "Adecuado", "Destacado"]
        bg_colors = ["EFEFEF", "FEF2F2", "FEFCE8", "F0FDF4", "EFF6FF"] # Gris, Rojo, Amarillo, Verde, Azul
        text_colors = ["000000", "991B1B", "854D0E", "166534", "1E40AF"]

        for i, (text, bg) in enumerate(zip(headers, bg_colors)):
            style_header_cell(t.cell(0,i), text, bg, "000000" if i==0 else "333333")

        # Filas
        puntaje_total = data.get('puntaje_total', 60)
        
        for row in tabla_datos:
            new_row = t.add_row()
            
            # Criterio y Peso
            peso = row.get('porcentaje', 0)
            pts_max = (puntaje_total * peso) / 100
            criterio_text = f"{limpiar_latex_para_word(row.get('criterio',''))}\n\n({peso}%) - Máx {pts_max:.1f} pts"
            new_row.cells[0].text = criterio_text
            
            # Niveles
            niveles = row.get('niveles', {})
            # Mapeo seguro de claves (a veces vienen en mayusculas o minusculas)
            niveles_safe = {k.lower(): v for k, v in niveles.items()}
            
            keys = ["insuficiente", "elemental", "adecuado", "destacado"]
            factors = [0.25, 0.5, 0.75, 1.0]
            
            for i, key in enumerate(keys):
                desc = niveles_safe.get(key, "")
                pts = pts_max * factors[i]
                new_row.cells[i+1].text = f"{limpiar_latex_para_word(desc)}\n\n({pts:.1f} pts)"

def renderizar_evaluacion(doc, data):
    """Motor Visual para Pruebas"""
    titulo = data.get('title') or data.get('titulo') or "Evaluación"
    desc = data.get('description', '')
    items = data.get('items', [])

    if isinstance(data, dict) and "student_version" in data and isinstance(data['student_version'], dict):
        titulo = data['student_version'].get('title', titulo)
        desc = data['student_version'].get('description', desc)
        items = data['student_version'].get('items', items)
        
    if not items and isinstance(data, dict) and "preguntas" in data:
        items = data.get("preguntas", [])

    h = doc.add_heading(limpiar_latex_para_word(titulo), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if desc:
        p = doc.add_paragraph(limpiar_latex_para_word(desc))
        p.italic = True
    
    doc.add_paragraph()
    
    # Datos alumno
    t = doc.add_table(rows=1, cols=2); t.style='Table Grid'
    t.cell(0,0).text = "Nombre: __________________________________"; t.cell(0,1).text = "Fecha: __________"
    doc.add_paragraph()

    for i, item in enumerate(items):
        p = doc.add_paragraph()
        run = p.add_run(f"{i+1}. {limpiar_latex_para_word(item.get('stem',''))}")
        run.bold = True
        p.add_run(f" ({item.get('points')} pts)")
        
        if item.get('type') == 'multiple_choice' and item.get('options'):
            for idx, opt in enumerate(item['options']):
                text_opt = opt.get('text', '') if isinstance(opt, dict) else str(opt)
                doc.add_paragraph(f"   {chr(97+idx)}) {limpiar_latex_para_word(text_opt)}")
        elif item.get('type') == 'true_false':
             doc.add_paragraph("   ___ V    ___ F")
        elif item.get('type') in ['short_answer', 'essay']:
             doc.add_paragraph("_" * 80)
             doc.add_paragraph("_" * 80)

    # --- RENDERIZADO DE PAUTA DOCENTE (SI EXISTE) ---
    guide = data.get('teacher_guide')
    if guide and 'answers' in guide:
        doc.add_page_break()
        h = doc.add_heading("Pauta de Corrección Docente", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("Este anexo es exclusivo para el profesor.").italic = True
        doc.add_paragraph()

        for ans in guide['answers']:
            # Intentar encontrar el item relacionado para contextualizar
            item_id = ans.get('related_item_id')
            
            p = doc.add_paragraph()
            p.add_run(f"Pregunta {item_id}: ").bold = True
            
            if ans.get('correct_answer'):
                p.add_run(f"{ans.get('correct_answer')}")
            
            if ans.get('explanation'):
                doc.add_paragraph(f"Justificación: {ans.get('explanation')}").style = 'Quote' # O 'Intense Quote'
            
            if ans.get('rubric'):
                r = ans['rubric']
                doc.add_paragraph("Rúbrica Sugerida:").bold = True
                tr = doc.add_table(rows=2, cols=3); tr.style = 'Table Grid'
                tr.cell(0,0).text = "LOGRADO"; tr.cell(0,1).text = "MEDIANAMENTE"; tr.cell(0,2).text = "NO LOGRADO"
                tr.cell(1,0).text = str(r.get('logrado',''))
                tr.cell(1,1).text = str(r.get('medianamente',''))
                tr.cell(1,2).text = str(r.get('no_logrado',''))
                doc.add_paragraph()

def renderizar_lectura_inteligente(doc, data: Dict[str, Any]):
    """Motor Visual para Lectura Inteligente"""
    is_profesor = data.get("tipo_documento", "estudiante") == "profesor"
    
    # Título principal
    h = doc.add_heading(f"Guía de Comprensión Lectora - {data.get('asignatura')}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Nivel: {data.get('nivel')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Objetivo: {data.get('oa')}\n").italic = True
    
    # 1. El Texto
    doc.add_heading("I. Lectura Base", level=2)
    p_texto = doc.add_paragraph(limpiar_latex_para_word(data.get("texto", "")))
    p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Espacio
    doc.add_paragraph()
    
    # 2. Las Preguntas
    doc.add_heading("II. Preguntas de Comprensión", level=2)
    preguntas = data.get("preguntas", [])
    
    for i, q in enumerate(preguntas):
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"{i + 1}. {limpiar_latex_para_word(q.get('pregunta', ''))}")
        run_q.bold = True
        
        # Opcional: Mostrar DOK en cursiva y pequeño (especialmente útil si es profesor)
        if is_profesor:
            p_q.add_run(f" [{q.get('nivel_taxonomico', '')}]").font.color.rgb = RGBColor(120, 120, 120)

        alts = q.get('alternativas', [])
        for j, alt in enumerate(alts):
            p_alt = doc.add_paragraph(f"   {chr(65+j)}) {limpiar_latex_para_word(alt)}")
            # Verificar si la alternativa es la correcta (búsqueda robusta)
            alt_correcta = str(q.get("respuesta_correcta", "")).strip()
            es_correcta = False
            if alt_correcta == alt:
                es_correcta = True
            elif alt_correcta.startswith(chr(65+j)) or alt_correcta.startswith(chr(97+j)):
                es_correcta = True
            elif f"opcion {chr(97+j)}" in alt_correcta.lower() or f"opción {chr(97+j)}" in alt_correcta.lower():
                 es_correcta = True

            # En la pauta del profesor, destacar la correcta
            if is_profesor and es_correcta:
                p_alt.runs[0].bold = True
                p_alt.runs[0].font.color.rgb = RGBColor(0, 128, 0) # Verde oscuro
        
        # En versión profesor, mostrar la justificación
        if is_profesor:
            p_just = doc.add_paragraph()
            r_just = p_just.add_run("Justificación Pedagógica: ")
            r_just.bold = True
            r_just.font.color.rgb = RGBColor(217, 154, 80) # Naranja (#d99a50)
            p_just.add_run(limpiar_latex_para_word(q.get("justificacion", "")))
        else:
             doc.add_paragraph()

# ==========================================
# 4. ENDPOINTS
# ==========================================

@router.post("/export/planificacion-docx")
async def export_planificacion_docx(req: PlanExportRequest):
    try:
        doc = Document()
        add_header_logo(doc, req.asignatura, req.nivel, "Planificación")
        renderizar_planificacion(doc, req.dict())
        file_stream = io.BytesIO(); doc.save(file_stream); file_stream.seek(0)
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=planificacion.docx"})
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/rubrica-docx")
async def export_rubric_docx(req: ExportRequest):
    try:
        doc = Document()
        add_header_logo(doc, req.asignatura, req.nivel, "Rúbrica de Evaluación")
        renderizar_rubrica(doc, req.dict())
        file_stream = io.BytesIO(); doc.save(file_stream); file_stream.seek(0)
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=rubrica.docx"})
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/evaluacion-docx")
async def export_assessment_docx(req: AssessmentExportRequest):
    try:
        doc = Document()
        add_header_logo(doc, req.subject, req.grade, "Evaluación")
        renderizar_evaluacion(doc, req.dict())
        file_stream = io.BytesIO(); doc.save(file_stream); file_stream.seek(0)
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=evaluacion.docx"})
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/elevador-docx")
async def export_elevator_docx(req: ElevatorExportRequest):
    try:
        doc = Document()
        add_header_logo(doc, req.subject, req.grade, "Elevador Cognitivo")
        doc.add_heading("Elevador Cognitivo", 1)
        doc.add_paragraph(f"Actividad Base: {req.activity}")
        doc.add_heading(f"Diagnóstico: {req.dok_actual}", 2)
        doc.add_paragraph(req.diagnostico)
        # (Se puede expandir la lógica visual aquí si es necesario)
        file_stream = io.BytesIO(); doc.save(file_stream); file_stream.seek(0)
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=elevador.docx"})
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/lectura-inteligente")
async def export_lectura_docx(req: LecturaInteligenteExportRequest):
    try:
        doc = Document()
        formato = "Pauta Docente" if req.tipo_documento == "profesor" else "Guía Estudiante"
        add_header_logo(doc, req.asignatura, req.nivel, formato)
        
        renderizar_lectura_inteligente(doc, req.dict())
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": f"attachment; filename=lectura_{req.tipo_documento}.docx"}
        )
    except Exception as e:
        print(f"Error DOCX: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- GENÉRICO (BIBLIOTECA) ---

@router.post("/export/prepare-generic")
async def prepare_generic(req: GenericExportRequest):
    req_id = str(uuid.uuid4())
    download_cache[req_id] = req
    return {"download_id": req_id}

@router.get("/export/download-generic/{req_id}/{filename}")
async def download_generic_get(req_id: str, filename: str):
    if req_id not in download_cache:
        raise HTTPException(status_code=404, detail="Descarga expirada o inválida")
    req = download_cache.pop(req_id)
    return await render_export_generic_docx(req)

def detectar_y_renderizar(doc, content: Any, titulo: str):
    """Dispatcher universal: detecta el tipo de documento y llama al motor de renderizado correcto."""
    if not isinstance(content, (dict, list)) or content is None:
        doc.add_paragraph(str(content))
        return

    # 1. PLANIFICACIÓN
    if isinstance(content, dict) and ("planificacion_clases" in content or ("clases" in content and isinstance(content.get('clases'), list))):
        if "titulo_unidad_creativo" in content: content["titulo_unidad"] = content["titulo_unidad_creativo"]
        if "estrategia_aprendizaje_sentencia" in content: content["estrategia"] = content["estrategia_aprendizaje_sentencia"]
        renderizar_planificacion(doc, content)

    # 2. RÚBRICA
    elif isinstance(content, dict) and "tabla" in content and isinstance(content.get('tabla'), list) and len(content['tabla']) > 0 and "criterio" in content['tabla'][0]:
        renderizar_rubrica(doc, content)

    # 3. EVALUACIÓN (todos los formatos)
    elif es_evaluacion(content):
        if isinstance(content, list):
            renderizar_evaluacion(doc, {"title": titulo, "items": content})
        else:
            renderizar_evaluacion(doc, content)

    # 4. LECTURA INTELIGENTE
    elif isinstance(content, dict) and "texto" in content and ("preguntas" in content or "oa" in content):
        renderizar_lectura_inteligente(doc, content)

    # 5. ELEVADOR COGNITIVO
    elif isinstance(content, dict) and ("escalera" in content or ("propuestas" in content and "dok_actual" in content)):
        doc.add_heading("Elevador Cognitivo", level=1)
        doc.add_paragraph(f"Actividad Base: {limpiar_latex_para_word(str(content.get('activity', content.get('actividad', ''))))}") 
        if content.get('dok_actual'): doc.add_paragraph(f"Nivel DOK Actual: {content.get('dok_actual')}").runs[0].bold = True
        if content.get('diagnostico'): doc.add_paragraph(limpiar_latex_para_word(str(content.get('diagnostico')))).italic = True
        escalera = content.get('escalera', [])
        if escalera:
            doc.add_heading("Escalera de Complejidad", level=2)
            for paso in escalera:
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"DOK {paso.get('dok_level', '?')}: ").bold = True
                p.add_run(limpiar_latex_para_word(str(paso.get('actividad', paso.get('activity', '')))))
        propuestas = content.get('propuestas', {})
        if propuestas:
            doc.add_heading("Propuestas de Mejora", level=2)
            for clave, val in propuestas.items():
                doc.add_paragraph(f"• {clave.upper()}: {limpiar_latex_para_word(str(val))}")

    # 6. NEE / ADECUACIÓN
    elif isinstance(content, dict) and ("estrategias" in content or "barrier" in content or "adaptaciones" in content):
        doc.add_heading("Plan de Adecuación Curricular (NEE)", level=1)
        if content.get('diagnosis'):
            p = doc.add_paragraph(); p.add_run("Diagnóstico: ").bold = True; p.add_run(content.get('diagnosis'))
        if content.get('barrier'):
            p = doc.add_paragraph(); p.add_run("Barrera Principal: ").bold = True; p.add_run(content.get('barrier'))
        estrategias = content.get('estrategias', content.get('adaptaciones', {}))
        if isinstance(estrategias, dict):
            for seccion, texto in estrategias.items():
                doc.add_heading(seccion.replace('_', ' ').title(), level=2)
                doc.add_paragraph(limpiar_latex_para_word(str(texto)))
        elif isinstance(estrategias, list):
            for est in estrategias:
                if isinstance(est, dict):
                    doc.add_paragraph(f"• {limpiar_latex_para_word(str(est.get('descripcion', est.get('text', str(est)))))}", style='List Bullet')

    # 7. AUDITORÍA / ANALIZADOR
    elif isinstance(content, dict) and "diagnostico_global" in content:
        doc.add_heading("Reporte de Auditoría Pedagógica", level=1)
        p = doc.add_paragraph(); p.add_run("Diagnóstico: ").bold = True; p.add_run(limpiar_latex_para_word(str(content.get('diagnostico_global', 'N/A'))))
        score = content.get('score_coherencia', 0)
        p2 = doc.add_paragraph()
        r = p2.add_run(f"Score de Coherencia: {score}%"); r.bold = True
        if isinstance(score, (int, float)) and score < 60: r.font.color.rgb = RGBColor(200, 0, 0)
        fortalezas = content.get('fortalezas', [])
        if fortalezas:
            doc.add_heading("Fortalezas", level=2)
            for f in fortalezas: doc.add_paragraph(f"✓ {limpiar_latex_para_word(str(f))}", style='List Bullet')
        brechas = content.get('brechas', content.get('oportunidades', []))
        if brechas:
            doc.add_heading("Oportunidades de Mejora", level=2)
            for b in brechas: doc.add_paragraph(f"→ {limpiar_latex_para_word(str(b))}", style='List Bullet')

    # FALLBACK INTELIGENTE: texto legible en lugar de JSON crudo
    else:
        doc.add_heading(titulo or "Recurso", level=1)
        if isinstance(content, dict):
            for key, val in content.items():
                if isinstance(val, str) and len(val) > 0:
                    p = doc.add_paragraph()
                    p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                    p.add_run(limpiar_latex_para_word(val))
                elif isinstance(val, list) and len(val) > 0:
                    doc.add_heading(key.replace('_', ' ').title(), level=2)
                    for elem in val:
                        if isinstance(elem, str): doc.add_paragraph(f"• {elem}", style='List Bullet')
                        elif isinstance(elem, dict) and 'text' in elem: doc.add_paragraph(f"• {elem['text']}", style='List Bullet')

async def render_export_generic_docx(req: GenericExportRequest):
    try:
        doc = Document()
        add_header_logo(doc, req.asignatura, req.nivel, "Documento Exportado")
        detectar_y_renderizar(doc, req.contenido, req.titulo_unidad)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"{req.titulo_unidad}.docx"
        headers = get_safe_headers(filename)
        headers["Access-Control-Expose-Headers"] = "Content-Disposition"
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers=headers
        )
    except Exception as e:
        print(f"Error export genérico: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/export/prepare-paquete")
async def prepare_paquete(req: PaqueteExportRequest):
    req_id = str(uuid.uuid4())
    download_cache[req_id] = req
    return {"download_id": req_id}

@router.get("/export/download-paquete/{req_id}/{filename}")
async def download_paquete_get(req_id: str, filename: str):
    if req_id not in download_cache:
        raise HTTPException(status_code=404, detail="Descarga expirada o inválida")
    req = download_cache.pop(req_id)
    return await render_export_paquete_docx(req)

async def render_export_paquete_docx(req: PaqueteExportRequest):
    try:
        doc = Document()
        
        add_header_logo(doc, "Paquete Didáctico", "Múltiples Niveles", "Documento Consolidado")
        
        for _ in range(5): doc.add_paragraph()
        
        main_title = doc.add_heading("PAQUETE DIDÁCTICO ÍNTEGRO", level=1)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in main_title.runs:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(27, 60, 115)
            
        subtitle = doc.add_paragraph("ProfeIC - Transformando la Educación")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(242, 174, 96)
        
        doc.add_paragraph(f"Total de recursos incluidos: {len(req.documentos)}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        for idx, item in enumerate(req.documentos):
            sep = doc.add_heading(f"--- DOCUMENTO {idx + 1}: {item.titulo_unidad} ---", level=2)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Asignatura: {item.asignatura} | Nivel: {item.nivel}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            detectar_y_renderizar(doc, item.contenido, item.titulo_unidad)

            if idx < len(req.documentos) - 1:
                doc.add_page_break()

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        headers = get_safe_headers("Paquete_Didactico_ProfeIC.docx")
        headers["Access-Control-Expose-Headers"] = "Content-Disposition"
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers=headers
        )
    except Exception as e:
        print("====== ERROR EN PAQUETE DOCX ======")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# EXECUTIVE REPORT RENDERER (PREMIUM)
# ==========================================
class ExecutiveDocxRequest(BaseModel):
    systemic_summary: str
    top_3_gaps: Union[List[str], List[Dict[str, Any]], List[Any]]
    recommended_training: Union[str, List[Dict[str, Any]]]
    rigor_audit: Optional[Dict[str, Any]] = None
    heatmap: Optional[Dict[str, float]] = None
    global_metrics: Optional[Dict[str, Any]] = None
    highlights: Optional[Dict[str, Any]] = None
    matriz: Optional[List[Dict[str, Any]]] = None

def get_color_for_score(score: float, invert: bool = False) -> str:
    """Returns a hex color based on the score (1-4 scale or percentages)."""
    # Percentage scale (0-100)
    if score > 4:
        if score >= 80: return "F0FDF4" if not invert else "FEF2F2" # Green / Red
        if score >= 60: return "FEFCE8" # Yellow
        return "FEF2F2" if not invert else "F0FDF4" # Red / Green
    
    # 1-4 scale
    if score >= 3.2: return "F0FDF4" if not invert else "FEF2F2"
    if score >= 2.6: return "FEFCE8"
    return "FEF2F2" if not invert else "F0FDF4"

def renderizar_reporte_ejecutivo(doc, data):
    # --- PÁGINA 1: PORTADA ---
    # Add spacing for vertical centering
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_heading("REPORTE EJECUTIVO INSTITUCIONAL", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(27, 60, 115) # #1B3C73
        
    subtitle = doc.add_paragraph("Análisis de Inteligencia Pedagógica")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(42, 89, 168) # #2A59A8
    
    for _ in range(3): doc.add_paragraph()
    
    if data.global_metrics:
        t_cover = doc.add_table(rows=1, cols=3)
        t_cover.style = 'Table Grid'
        t_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        c1, c2, c3 = t_cover.rows[0].cells
        style_header_cell(c1, "DOCENTES EN PLANTA", "1B3C73", "FFFFFF")
        c1.add_paragraph(str(data.global_metrics.get('total_teachers', 0))).alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1.paragraphs[1].runs[0].font.size = Pt(18); c1.paragraphs[1].runs[0].bold = True
        
        style_header_cell(c2, "COBERTURA", "A1C969", "FFFFFF")
        c2.add_paragraph(f"{data.global_metrics.get('coverage_percent', 0)}%").alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.paragraphs[1].runs[0].font.size = Pt(18); c2.paragraphs[1].runs[0].bold = True
        
        style_header_cell(c3, "CICLOS CERRADOS", "C87533", "FFFFFF")
        c3.add_paragraph(str(data.global_metrics.get('total_completed', 0))).alignment = WD_ALIGN_PARAGRAPH.CENTER
        c3.paragraphs[1].runs[0].font.size = Pt(18); c3.paragraphs[1].runs[0].bold = True
        
    doc.add_page_break()

    # --- PÁGINA 2: DIAGNÓSTICO Y HEATMAP ---
    doc.add_heading("I. Diagnóstico Sistémico", level=2)
    p = doc.add_paragraph(data.systemic_summary)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- PÁGINA 3: ESTRUCTURA DE ACOMPAÑAMIENTO Y DEPARTAMENTOS ---
    if data.global_metrics and "structural" in data.global_metrics:
        doc.add_page_break()
        doc.add_heading("II. Estructura de Acompañamiento", level=2)
        doc.add_paragraph("Relación de docentes acompañados frente al equipo de evaluadores activos en este período.")
        
        struct_data = data.global_metrics["structural"]
        
        t_struct = doc.add_table(rows=1, cols=2)
        t_struct.style = 'Table Grid'
        t_struct.autofit = False
        t_struct.columns[0].width = Inches(2.75)
        t_struct.columns[1].width = Inches(2.75)
        
        c1, c2 = t_struct.rows[0].cells
        style_header_cell(c1, "EVALUADORES ACTIVOS", "1B3C73", "FFFFFF")
        style_header_cell(c2, "DOCENTES ACOMPAÑADOS", "1B3C73", "FFFFFF")
        
        row_struct = t_struct.add_row().cells
        row_struct[0].text = str(struct_data.get("unique_observers", 0))
        row_struct[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_struct[1].text = str(struct_data.get("unique_teachers", 0))
        row_struct[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        depts = struct_data.get("departments", {})
        if depts:
            doc.add_heading("Distribución por Departamentos", level=3)
            doc.add_paragraph("Volumen de observaciones completadas según el departamento al que pertenece el docente.")
            
            t_dept = doc.add_table(rows=1, cols=2)
            t_dept.style = 'Table Grid'
            t_dept.autofit = False
            t_dept.columns[0].width = Inches(4.0)
            t_dept.columns[1].width = Inches(1.5)
            
            hdr_dept = t_dept.rows[0].cells
            style_header_cell(hdr_dept[0], "Departamento", "A1C969", "FFFFFF")
            style_header_cell(hdr_dept[1], "Nº Obs.", "A1C969", "FFFFFF")
            
            # Sort by count descending
            for d_name, count in sorted(depts.items(), key=lambda x: x[1], reverse=True):
                r_dept = t_dept.add_row().cells
                r_dept[0].text = str(d_name).title()
                r_dept[1].text = str(count)
                r_dept[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.heatmap:
        doc.add_page_break()
        doc.add_heading("III. Pulso Cuantitativo (Heatmap)", level=2)
        doc.add_paragraph("El siguiente mapa de calor muestra el promedio institucional (en escala 1 a 4) para las distintas dimensiones observadas.")
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Inches(4.0)
        table.columns[1].width = Inches(1.5)
        
        hdr_cells = table.rows[0].cells
        style_header_cell(hdr_cells[0], "Dimensión Pedagógica", "1B3C73", "FFFFFF")
        style_header_cell(hdr_cells[1], "Puntaje (1-4)", "1B3C73", "FFFFFF")
        
        # Sort heatmap by value (descending) to show strengths first
        sorted_heatmap = sorted(data.heatmap.items(), key=lambda x: x[1], reverse=True)
        
        for k, v in sorted_heatmap:
            row_cells = table.add_row().cells
            row_cells[0].text = k.replace("_", " ").title()
            row_cells[0].paragraphs[0].runs[0].bold = True
            
            row_cells[1].text = f"{v:.1f}"
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].runs[0].bold = True
            
            # Apply color to the score cell
            bg_color = get_color_for_score(v)
            set_cell_background(row_cells[1], bg_color)
            
    doc.add_page_break()

    # --- PÁGINA X: MATRIZ DE ACCIÓN ---
    doc.add_page_break()
    doc.add_heading("IV. Matriz de Acción Estratégica", level=2)
    
    # Brechas
    p_brechas = doc.add_paragraph()
    p_brechas.add_run("Principales Brechas Detectadas").bold = True
    for i, gap in enumerate(data.top_3_gaps):
        gap_str = gap.get("descripcion", gap.get("gap", str(gap))) if isinstance(gap, dict) else str(gap)
        doc.add_paragraph(f"{gap_str}", style='List Number')
        
    doc.add_paragraph()
    
    # Matriz de Capacitación
    p_cap = doc.add_paragraph()
    p_cap.add_run("Plan de Intervención Recomendado").bold = True
    
    training_list = data.recommended_training
    
    if isinstance(training_list, list) and len(training_list) > 0:
        t_action = doc.add_table(rows=1, cols=4)
        t_action.style = 'Table Grid'
        t_action.autofit = False
        
        # Set column widths roughly
        t_action.columns[0].width = Inches(1.5)
        t_action.columns[1].width = Inches(1.5)
        t_action.columns[2].width = Inches(2.2)
        t_action.columns[3].width = Inches(1.3)
        
        style_header_cell(t_action.cell(0,0), "Iniciativa (Foco)", "C87533", "FFFFFF")
        style_header_cell(t_action.cell(0,1), "Objetivo Esperado", "C87533", "FFFFFF")
        style_header_cell(t_action.cell(0,2), "Metodología de Implem.", "C87533", "FFFFFF")
        style_header_cell(t_action.cell(0,3), "KPI de Éxito", "C87533", "FFFFFF")
        
        for item in training_list:
            if not isinstance(item, dict): continue
            row = t_action.add_row().cells
            
            row[0].text = item.get("foco", "")
            row[0].paragraphs[0].runs[0].bold = True
            
            row[1].text = item.get("objetivo", "")
            row[2].text = item.get("metodologia", "")
            
            row[3].text = item.get("kpi", "")
            row[3].paragraphs[0].runs[0].italic = True
    else:
        p_train = doc.add_paragraph(str(training_list))
        if p_train.runs:
            p_train.runs[0].italic = True

    # --- PÁGINA X: ANEXOS Y CUADRO DE HONOR ---
    if data.highlights or data.rigor_audit:
        doc.add_page_break()
        doc.add_heading("V. Anexos y Destacados", level=2)

    if data.highlights:
        # Cuadro de Honor
        t_honor = doc.add_table(rows=1, cols=2)
        t_honor.autofit = False
        t_honor.columns[0].width = Inches(3.0)
        t_honor.columns[1].width = Inches(3.0)
        c_left, c_right = t_honor.rows[0].cells
        
        # Left: Top Teachers
        style_header_cell(c_left, "CUADRO DE HONOR: DOCENTES CON MAYOR MEJORA", "A1C969", "FFFFFF")
        top_t = data.highlights.get('top_teachers', [])
        if top_t:
            for t in top_t:
                c_left.add_paragraph(f"• {t.get('name')} (+{t.get('score')} pts)")
        else:
            p_empty1 = c_left.add_paragraph("No hay datos en este período.")
            if p_empty1.runs: p_empty1.runs[0].italic = True
            
        # Right: Top Observers
        style_header_cell(c_right, "LIDERAZGO PEDAGÓGICO: ACOMPAÑANTES DESTACADOS", "2A59A8", "FFFFFF")
        top_o = data.highlights.get('top_observers', [])
        if top_o:
            for o in top_o:
                c_right.add_paragraph(f"• {o.get('name')} (KPI: {o.get('kpi_score')})")
        else:
            p_empty2 = c_right.add_paragraph("No hay datos en este período.")
            if p_empty2.runs: p_empty2.runs[0].italic = True
            
        doc.add_paragraph()

    # Auditoría Final
    if data.rigor_audit:
        p_rigor = doc.add_paragraph()
        run = p_rigor.add_run("Auditoría de Rigor del Proceso de Evaluación:")
        run.bold = True
        
        t_rigor = doc.add_table(rows=1, cols=3)
        t_rigor.style = 'Table Grid'
        
        c1, c2, c3 = t_rigor.rows[0].cells
        style_header_cell(c1, "Índice de Profundidad", "EFEFEF", "000000")
        c1.add_paragraph(f"{data.rigor_audit.get('depth_index', 0)}%").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        style_header_cell(c2, "Estado (IA)", "EFEFEF", "000000")
        alert = data.rigor_audit.get('alert', False)
        status_text = "ALERTA" if alert else "ÓPTIMO"
        c2.add_paragraph(status_text).alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(c2, "FEF2F2" if alert else "F0FDF4")
        
        style_header_cell(c3, "Muestra Analizada", "EFEFEF", "000000")
        c3.add_paragraph(f"{data.rigor_audit.get('sample_size', 0)} obs.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- NÚCLEO: MATRIZ DE ACOMPAÑAMIENTO (ANEXO) ---
    if data.matriz and len(data.matriz) > 0:
        doc.add_page_break()
        doc.add_heading("VI. Anexo: Registro Histórico de Acompañamiento", level=2)
        doc.add_paragraph("Detalle de las observaciones realizadas durante el período analizado.")
        
        t_matriz = doc.add_table(rows=1, cols=4)
        t_matriz.style = 'Table Grid'
        t_matriz.autofit = False
        t_matriz.columns[0].width = Inches(1.0) # Fecha
        t_matriz.columns[1].width = Inches(2.2) # Docente
        t_matriz.columns[2].width = Inches(2.2) # Observador
        t_matriz.columns[3].width = Inches(1.1) # Estado
        
        cells = t_matriz.rows[0].cells
        style_header_cell(cells[0], "Fecha", "1B3C73", "FFFFFF")
        style_header_cell(cells[1], "Docente Observado", "1B3C73", "FFFFFF")
        style_header_cell(cells[2], "Acompañante", "1B3C73", "FFFFFF")
        style_header_cell(cells[3], "Estado", "1B3C73", "FFFFFF")
        
        # Sort data by date descending just in case
        try:
            sorted_matriz = sorted(data.matriz, key=lambda x: x.get('date', ''), reverse=True)
        except Exception:
            sorted_matriz = data.matriz
            
        for m in sorted_matriz:
            row = t_matriz.add_row().cells
            # Parse Date logic (safely)
            date_str = str(m.get('date', '')).split('T')[0] if m.get('date') else 'N/A'
            row[0].text = date_str
            row[1].text = str(m.get('teacher_name', ''))
            row[2].text = str(m.get('observer_name', ''))
            
            status = m.get('status', '')
            status_text = "Completado" if status == 'completed' else "En Proceso" if status == 'in_progress' else "Planificado"
            row[3].text = status_text
            
            # Color status
            if status == 'completed': set_cell_background(row[3], "F0FDF4") # Verde
            elif status == 'in_progress': set_cell_background(row[3], "EFF6FF") # Azul
            else: set_cell_background(row[3], "FEFCE8") # Amarillo

@router.post("/export/executive-docx")
async def export_executive_docx(req: ExecutiveDocxRequest):
    try:
        doc = Document()
        # Modificar logo original solo si es necesario, o lo metemos con la portada arriba
        add_header_logo(doc, "Dirección Académica", "Institucional", "Reporte Confidencial")
        renderizar_reporte_ejecutivo(doc, req)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": "attachment; filename=reporte_ejecutivo.docx"}
        )
    except Exception as e:
        print(f"Error DOCX: {e}")
        raise HTTPException(status_code=500, detail=str(e))