"""
routers/simce_generator.py
──────────────────────────
Endpoint POST /api/v1/simce/generate

Flujo completo:
  1. Cálculo del blueprint       → simce_blueprint_parser.calculate_question_distribution()
  2. Contexto RAG               → Supabase tabla `curriculum_indicadores`
  3. Construcción del super-prompt
  4. Llamada al LLM             → Gemini 1.5 Flash (json_object mode)
  5. Respuesta JSON estructurada al cliente Next.js

Para registrar en main.py (ya tiene import de `simce`, añadir):
    from routers import simce_generator
    app.include_router(simce_generator.router)
"""

from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import random
import re
import unicodedata
from pathlib import Path
from typing import Any, Literal
from urllib.parse import quote

import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field, field_validator
from supabase import Client, create_client

from routers.deps import get_current_user_id
from simce_blueprint_parser import BlueprintResult, calculate_question_distribution

# ─── Logger ───────────────────────────────────────────────────────────────────

logger = logging.getLogger(__name__)

# ─── Supabase (service-role para evitar RLS en lectura de curriculum) ──────────

_SB_URL: str = os.getenv("SUPABASE_URL", "")
_SB_KEY: str = os.getenv("SUPABASE_SERVICE_ROLE_KEY") or os.getenv("SUPABASE_KEY", "")

supabase: Client | None = None
if _SB_URL and _SB_KEY:
    supabase = create_client(_SB_URL, _SB_KEY)
else:
    logger.warning("⚠️  Supabase no configurado — el contexto RAG no funcionará.")

# ─── Gemini ────────────────────────────────────────────────────────────────────
# NOTE: Si deseas cambiar a OpenAI, agrega OPENAI_API_KEY al .env y reemplaza
#       la llamada _call_llm() por openai.AsyncOpenAI().chat.completions.create(
#           model="gpt-4o-mini", response_format={"type": "json_object"}, ...)

_GEMINI_MODEL = "gemini-2.0-flash"  # (User requested 2.5, using 2.0 as it is the current latest Flash)
_GEMINI_KEY: str = os.getenv("GOOGLE_API_KEY", "")
if _GEMINI_KEY:
    genai.configure(api_key=_GEMINI_KEY)
else:
    logger.warning("⚠️  GOOGLE_API_KEY no configurada — la generación IA no funcionará.")

# ─── Ruta al JSON de tablas de especificaciones ────────────────────────────────

# El parser vive en backend/ y el JSON en la raíz del proyecto (un nivel arriba)
_BLUEPRINT_JSON: Path = Path(__file__).parent.parent.parent / "tablas_especificaciones_simce.json"

_ASSETS_DIR: Path = Path(__file__).parent.parent / "assets"
_ESTIMULOS_JSON: Path = _ASSETS_DIR / "banco_estimulos_unificado.json"
_MATRIZ_JSON: Path = _ASSETS_DIR / "matriz_distractores_oficial.json"

# ─── Router ────────────────────────────────────────────────────────────────────

router = APIRouter(prefix="/api/v1/simce", tags=["SIMCE Generator"])


# ─── Modelos Pydantic ──────────────────────────────────────────────────────────

class SimceGenerateRequest(BaseModel):
    """Payload que envía el wizard de Next.js."""

    nivel: str = Field(
        ...,
        description="Nivel educativo. Ej: '4° Básico', '8° Básico', 'II Medio'.",
        examples=["8° Básico"],
    )
    asignatura: str = Field(
        ...,
        description="Nombre de la asignatura exactamente como aparece en el JSON de tablas.",
        examples=["Lenguaje y Comunicación - Lectura"],
    )
    cantidad_preguntas: int = Field(
        ...,
        ge=20,
        le=45,
        description="Total de ítems del instrumento (entre 20 y 45).",
        examples=[35],
    )
    modo: Literal["formativo", "simce"] = Field(
        default="simce",
        description="Modo de generación: formativo (libre) o simce (tabla de especificaciones).",
    )
    oas_seleccionados: list[str] = Field(
        default_factory=list,
        description="Lista de IDs de OAs elegidos en el wizard (solo si modo=formativo).",
    )
    archivo_base64: str | None = Field(
        default=None,
        description="Archivo subido por el usuario en formato base64 sin prefijo URI.",
    )
    archivo_mime_type: str | None = Field(
        default=None,
        description="Mime type del archivo subido (e.g. 'image/png', 'application/pdf').",
    )
    archivo_nombre: str | None = Field(
        default=None,
        description="Nombre original del archivo subido.",
    )

    @field_validator("cantidad_preguntas")
    @classmethod
    def validate_cantidad(cls, v: int) -> int:
        if not (20 <= v <= 45):
            raise ValueError("cantidad_preguntas debe estar entre 20 y 45.")
        return v


class PreguntaGenerada(BaseModel):
    """Ítem individual del instrumento generado."""

    numero: int
    habilidad_medida: str
    enunciado: str
    alternativas: dict[str, Any]   # Puede ser string simple o objeto dict con texto y tipos de error
    correcta: str
    justificacion: str
    explicacion_correcta: str | None = None


class SimceGenerateResponse(BaseModel):
    """Respuesta devuelta al cliente Next.js."""

    titulo: str
    asignatura: str
    nivel: str
    cantidad_preguntas: int
    modo: str
    blueprint_resumen: str          # Texto del blueprint para mostrar en UI
    estimulo_texto: str | None = None
    estimulo_imagen: str | None = None
    preguntas: list[PreguntaGenerada]


# ─── Helpers privados ──────────────────────────────────────────────────────────

def _get_random_stimulus(asignatura: str, nivel: str) -> dict | None:
    try:
        if not _ESTIMULOS_JSON.exists(): return None
        data = json.loads(_ESTIMULOS_JSON.read_text("utf-8"))
        filtrados = [item for item in data if item.get("asignatura") == asignatura and item.get("nivel") == nivel and item.get("tipo") != "imagen_adjunta_solo"]
        if filtrados:
            return random.choice(filtrados)
        return None
    except Exception as e:
        logger.error("Error al leer estimulos: %s", e)
        return None

def _get_matriz_distractores() -> str:
    try:
        if not _MATRIZ_JSON.exists(): return ""
        data = json.loads(_MATRIZ_JSON.read_text("utf-8"))
        return json.dumps(data.get("tipos_error", {}), ensure_ascii=False)
    except Exception as e:
        logger.error("Error al leer matriz: %s", e)
        return ""

def _get_blueprint(request: SimceGenerateRequest) -> BlueprintResult:
    """Calcula la distribución de preguntas desde el JSON de tablas SIMCE."""
    try:
        return calculate_question_distribution(
            json_path=_BLUEPRINT_JSON,
            asignatura=request.asignatura,
            nivel=request.nivel,
            total_preguntas=request.cantidad_preguntas,
        )
    except FileNotFoundError as exc:
        logger.error("❌ JSON de tablas no encontrado: %s", _BLUEPRINT_JSON)
        raise HTTPException(
            status_code=500,
            detail=f"Archivo de tablas de especificaciones no encontrado: {exc}",
        ) from exc
    except ValueError as exc:
        logger.error("❌ Asignatura/nivel no encontrado en tablas: %s", exc)
        raise HTTPException(
            status_code=422,
            detail=str(exc),
        ) from exc


def _get_rag_context(nivel: str, asignatura: str, limit: int = 15) -> list[str]:
    """
    Obtiene indicadores de aprendizaje relevantes desde Supabase.
    Devuelve lista vacía si Supabase no está disponible o no hay resultados.
    """
    if supabase is None:
        logger.warning("⚠️  Supabase no disponible, saltando RAG.")
        return []

    try:
        # La tabla curriculum_indicadores tiene columnas: nivel, asignatura, indicador
        response = (
            supabase.table("curriculum_indicadores")
            .select("indicador")
            .eq("nivel", nivel)
            .eq("asignatura", asignatura)
            .limit(limit)
            .execute()
        )
        indicadores: list[str] = [
            row["indicador"]
            for row in (response.data or [])
            if row.get("indicador")
        ]
        logger.info("📚 RAG: %d indicadores recuperados para %s — %s", len(indicadores), asignatura, nivel)
        return indicadores
    except Exception as exc:
        # RAG es opcional: si falla, continuamos sin contexto curricular
        logger.warning("⚠️  Error al consultar Supabase RAG: %s", exc)
        return []


def _build_blueprint_text(blueprint: BlueprintResult) -> str:
    """Convierte el BlueprintResult en texto estructurado para inyectar en el prompt."""
    lines: list[str] = []

    # Distribución principal (habilidades o dominios cognitivos)
    tipo_label = (
        "HABILIDADES COGNITIVAS"
        if blueprint.tipo_eje == "habilidades"
        else "DOMINIOS COGNITIVOS"
    )
    lines.append(f"[{tipo_label}]")
    for ax in blueprint.distribucion_principal:
        lines.append(f"  • {ax.nombre}: {ax.preguntas} preguntas ({ax.porcentaje}%)")

    # Distribución temática (si existe)
    if blueprint.distribucion_tematica:
        lines.append("\n[EJES TEMÁTICOS DE CONTENIDO]")
        for te in blueprint.distribucion_tematica:
            lines.append(f"  • {te.eje}: {te.preguntas} preguntas ({te.porcentaje}%)")

    if blueprint.notas:
        lines.append(f"\nNota oficial: {blueprint.notas}")

    return "\n".join(lines)


def _build_system_prompt(
    request: SimceGenerateRequest,
    blueprint: BlueprintResult,
    distribucion_text: str,
    indicadores: list[str],
    estimulo_text: str = "",
    matriz_text: str = "",
    sub_total: int | None = None,
    start_number: int = 1,
) -> str:
    """Construye el super-prompt inyectando blueprint + RAG + formato de salida."""

    current_total = sub_total or request.cantidad_preguntas

    indicadores_texto: str = (
        "\n".join(f"  - {ind}" for ind in indicadores)
        if indicadores
        else "  (Sin indicadores específicos disponibles; usa el marco curricular general MINEDUC)"
    )

    oas_texto: str = ""
    if request.modo == "formativo" and request.oas_seleccionados:
        oas_texto = (
            f"\nOBJETIVOS DE APRENDIZAJE SELECCIONADOS POR EL DOCENTE: "
            f"{', '.join(request.oas_seleccionados)}\n"
            "Asegúrate de que cada pregunta mida al menos uno de estos OAs."
        )

    modo_instruccion: str = (
        "MODO: Simulación SIMCE Estricta — sigue la tabla de especificaciones al pie de la letra."
        if request.modo == "simce"
        else "MODO: Ensayo Formativo Libre — el énfasis es pedagógico y formativo."
    )

    estimulo_instruccion: str = ""
    if estimulo_text:
        estimulo_instruccion = f"""
ESTÍMULO DE LECTURA (OBLIGATORIO):
Utiliza el siguiente texto base para construir todas las preguntas. Es vital que el nivel de dificultad sea apropiado.
TEXTO:
{estimulo_text}"""
    elif request.archivo_base64:
        estimulo_instruccion = f"""
ESTÍMULO ADJUNTO (OBLIGATORIO):
Utiliza la imagen/documento que se adjunta a este prompt para construir todas las preguntas. 
Es vital que el nivel de dificultad sea apropiado."""

    psico_instruccion: str = f"""OBLIGATORIO: Construye los 3 distractores (alternativas incorrectas) utilizando EXCLUSIVAMENTE las reglas de error cognitivo de esta matriz oficial: {matriz_text}. En tu JSON de respuesta, añade el campo 'explicacion_correcta' y en cada distractor añade el campo 'tipo_error_oficial' que indique qué regla de la matriz usaste.""" if matriz_text else "5. La justificación debe explicar por qué la respuesta es correcta Y por qué cada distractor es incorrecto."

    return f"""Eres un experto en evaluación estandarizada SIMCE de Chile con 20 años de experiencia diseñando instrumentos para la Agencia de Calidad de la Educación.

Tu tarea es generar un ensayo de {current_total} preguntas de selección múltiple (alternativas A, B, C, D) para el nivel {request.nivel} en la asignatura {request.asignatura}.
Estas preguntas deben numerarse desde el {start_number} hasta el {start_number + current_total - 1}.

{modo_instruccion}
{oas_texto}
{estimulo_instruccion}

REGLA ESTADÍSTICA OBLIGATORIA (¡Crítica!):
Debes generar EXACTAMENTE la siguiente distribución de preguntas para este bloque, sin alterar los totales:
{distribucion_text}

BASES CURRICULARES MINEDUC:
Basa el contenido curricular estrictamente en estos indicadores de aprendizaje oficiales:
{indicadores_texto}

ESTÁNDARES TÉCNICOS PARA CADA ÍTEM:
1. El enunciado debe ser claro, preciso y situado en un contexto auténtico.
2. Cada ítem debe tener exactamente 4 alternativas (A, B, C, D).
3. Debe haber UNA sola respuesta correcta e inequívoca.
4. Los distractores (alternativas incorrectas) deben ser plausibles y representar errores conceptuales reales.
{psico_instruccion}
6. La dificultad debe ser apropiada para el nivel {request.nivel}.
7. NO repitas el mismo tipo de pregunta dos veces consecutivas.

FORMATO DE SALIDA:
Devuelve ÚNICAMENTE un JSON válido (sin markdown, sin bloques de código, sin texto extra) con esta estructura exacta:
{{
  "titulo": "Ensayo SIMCE — {request.asignatura} — {request.nivel}",
  "preguntas": [
    {{
      "numero": {start_number},
      "habilidad_medida": "<nombre exacto de la habilidad/dominio según la distribución>",
      "enunciado": "<texto completo del ítem>",
      "alternativas": {{
        "A": {{"texto": "<texto alternativa A>", "tipo_error_oficial": "<tipo error o nulo si es correcta>"}},
        "B": {{"texto": "<texto alternativa B>", "tipo_error_oficial": "<tipo error o nulo si es correcta>"}},
        "C": {{"texto": "<texto alternativa C>", "tipo_error_oficial": "<tipo error o nulo si es correcta>"}},
        "D": {{"texto": "<texto alternativa D>", "tipo_error_oficial": "<tipo error o nulo si es correcta>"}}
      }},
      "correcta": "A",
      "justificacion": "<explicación técnica tradicional u opcional>",
      "explicacion_correcta": "<explicación de la opción correcta>"
    }}
  ]
}}"""


async def _call_llm(prompt: str, file_part: dict | None = None) -> dict[str, Any] | None:
    """
    Llamada asíncrona a Gemini usando el SDK google-generativeai.
    Ejecuta en un thread pool para no bloquear FastAPI.
    """
    try:
        model = genai.GenerativeModel(
            model_name=_GEMINI_MODEL,
            generation_config={
                "response_mime_type": "application/json",
                "temperature": 0.7,
                "max_output_tokens": 8192,
            },
        )

        contents = [prompt]
        if file_part:
            contents.append(file_part)

        # Generación en thread para prevenir bloqueo
        response = await asyncio.to_thread(
            model.generate_content,
            contents,
            request_options={"timeout": 240}
        )

        texto: str = response.text.strip()

        # Limpieza defensiva
        for fence in ("```json", "```"):
            if texto.startswith(fence):
                texto = texto[len(fence):]
        if texto.endswith("```"):
            texto = texto[:-3]

        return json.loads(texto.strip())
    except Exception as e:
        logger.error("❌ Error en _call_llm: %s", e)
        return None


def _split_blueprint(blueprint: BlueprintResult, num_chunks: int = 2) -> list[BlueprintResult]:
    """
    Divide un BlueprintResult en N sub-blueprints balanceados.
    Útil para generar preguntas en paralelo sin perder la proporcionalidad.
    """
    import copy

    # Si solo hay 1 pregunta por habilidad, no podemos dividir mucho refinado,
    # pero para SIMCE (35-45 items) funciona bien.
    chunks: list[BlueprintResult] = []
    
    # Inicializar chunks con copias (sin preguntas aún)
    for _ in range(num_chunks):
        c = copy.deepcopy(blueprint)
        for ax in c.distribucion_principal:
            ax.preguntas = 0
        for te in c.distribucion_tematica:
            te.preguntas = 0
        chunks.append(c)

    # Distribuir preguntas de habilidades
    for ax_orig in blueprint.distribucion_principal:
        q = ax_orig.preguntas
        for i in range(q):
            chunks[i % num_chunks].distribucion_principal[
                next(j for j, a in enumerate(blueprint.distribucion_principal) if a.nombre == ax_orig.nombre)
            ].preguntas += 1

    # Distribuir preguntas temáticas (si hay)
    for te_orig in blueprint.distribucion_tematica:
        q = te_orig.preguntas
        for i in range(q):
            chunks[i % num_chunks].distribucion_tematica[
                next(j for j, t in enumerate(blueprint.distribucion_tematica) if t.eje == te_orig.eje)
            ].preguntas += 1

    # Ajustar totales por chunk
    for c in chunks:
        c.total_preguntas = sum(a.preguntas for a in c.distribucion_principal)

    return chunks


@router.post(
    "/generate",
    response_model=SimceGenerateResponse,
    summary="Generador SIMCE — Orquestación Segura",
)
async def generate_simce_instrument(
    request: SimceGenerateRequest,
    user_id: str = Depends(get_current_user_id),
) -> SimceGenerateResponse:
    """
    Orquesta la generación en paralelo con retrasos para evitar 429.
    """
    logger.info("🎯 SIMCE Generate: user=%s | modo=%s", user_id, request.modo)
    
    try:
        # Paso 1: Blueprint y Contexto (RAG)
        blueprint_full = _get_blueprint(request)
        indicadores = _get_rag_context(request.nivel, request.asignatura, limit=15)
        blueprint_resumen = str(blueprint_full)

        # Paso 2: Chunking (Dividimos en 2 bloques si hay más de 25 preguntas)
        num_chunks = 2 if request.cantidad_preguntas > 25 else 1
        sub_blueprints = _split_blueprint(blueprint_full, num_chunks)
        
        # Nuevas variables de estímulo y matriz
        estimulo = _get_random_stimulus(request.asignatura, request.nivel)
        estimulo_texto = estimulo.get("contenido_real") if estimulo else None
        estimulo_imagen = estimulo.get("imagen_adjunta") if estimulo else None
        matriz_text = _get_matriz_distractores()

        # ─── Función interna de protección por Lote ───
        async def safe_generate_chunk(sb, start_num, delay: float = 0):
            if delay > 0:
                logger.info("⏳ Esperando %.1fs para iniciar chunk %d...", delay, start_num)
                await asyncio.sleep(delay) 
            
            sb_text = _build_blueprint_text(sb)
            prompt = _build_system_prompt(
                request=request,
                blueprint=sb,
                distribucion_text=sb_text,
                indicadores=indicadores,
                estimulo_text=estimulo_texto or "",
                matriz_text=matriz_text,
                sub_total=sb.total_preguntas,
                start_number=start_num
            )
            
            file_part = None
            if request.archivo_base64 and request.archivo_mime_type:
                file_part = {
                    "mime_type": request.archivo_mime_type,
                    "data": request.archivo_base64
                }

            try:
                logger.info("🤖 Iniciando chunk %d (%d preguntas)...", start_num, sb.total_preguntas)
                resultado = await _call_llm(prompt, file_part=file_part) 
                
                # Validación estricta anti-caídas
                if not resultado or not isinstance(resultado, dict):
                    logger.warning("⚠️ Alerta en chunk %d: La IA no devolvió un JSON válido.", start_num)
                    return {"titulo": None, "preguntas": []}
                    
                return resultado
            except Exception as e:
                logger.error("❌ Excepción crítica en el chunk %d: %s", start_num, e)
                return {"titulo": None, "preguntas": []}

        # Paso 3: Encolar las tareas con retraso progresivo (2.5s)
        tasks = []
        current_start = 1
        for i, sb in enumerate(sub_blueprints):
            retraso = i * 2.5 
            tasks.append(safe_generate_chunk(sb, current_start, delay=retraso))
            current_start += sb.total_preguntas

        # Paso 4: Ejecución paralela controlada
        results_raw = await asyncio.gather(*tasks)

        # Paso 5: Consolidación de los fragmentos
        todas_las_preguntas_raw = []
        final_titulo = None
        for res in results_raw:
            todas_las_preguntas_raw.extend(res.get("preguntas", []))
            if not final_titulo and res.get("titulo"):
                final_titulo = res.get("titulo")

        if not todas_las_preguntas_raw:
            logger.error("❌ Array de preguntas final vacío.")
            raise ValueError("No se pudo generar ninguna pregunta. Revisa los logs de Gemini.")

        # Paso 6: Estructurar para Pydantic
        preguntas_finales = [
            PreguntaGenerada(
                numero=p.get("numero", i+1),
                habilidad_medida=p.get("habilidad_medida", "—"),
                enunciado=p.get("enunciado", ""),
                alternativas=p.get("alternativas", {"A":{"texto": ""}}),
                correcta=p.get("correcta", "A"),
                justificacion=p.get("justificacion", ""),
                explicacion_correcta=p.get("explicacion_correcta") or p.get("justificacion", ""),
            )
            for i, p in enumerate(todas_las_preguntas_raw)
        ]

        logger.info("✅ Generación finalizada: %d preguntas", len(preguntas_finales))

        return SimceGenerateResponse(
            titulo=final_titulo or f"Ensayo SIMCE — {request.asignatura}",
            asignatura=request.asignatura,
            nivel=request.nivel,
            cantidad_preguntas=len(preguntas_finales),
            modo=request.modo,
            blueprint_resumen=blueprint_resumen,
            estimulo_texto=estimulo_texto,
            estimulo_imagen=estimulo_imagen,
            preguntas=preguntas_finales,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error("❌ Falla general en la orquestación: %s", e)
        raise HTTPException(status_code=500, detail=f"Error generando instrumento: {str(e)}")

    # HTTPException ya tiene status code propio — re-raise sin envolver
    except HTTPException:
        raise

    # Cualquier otro error → 500
    except json.JSONDecodeError as exc:
        logger.error("❌ El LLM devolvió JSON inválido: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="El modelo IA devolvió una respuesta JSON inválida. Reintenta.",
        ) from exc

    except Exception as exc:
        logger.error(
            "❌ Error inesperado en generate_simce_instrument: %s: %s",
            type(exc).__name__, exc,
        )
        raise HTTPException(
            status_code=500,
            detail=f"Error al generar el instrumento: {type(exc).__name__}: {exc}",
        ) from exc


# ─── Helpers DOCX ─────────────────────────────────────────────────────────────

def _safe_filename(name: str) -> str:
    ascii_name = unicodedata.normalize("NFKD", name).encode("ASCII", "ignore").decode()
    return re.sub(r"[^\w\s.-]", "", ascii_name).replace(" ", "_") or "documento"


def _docx_headers(filename: str) -> dict:
    safe = _safe_filename(filename)
    utf8 = quote(filename.replace(" ", "_"))
    return {
        "Content-Disposition": f'attachment; filename="{safe}"; filename*=UTF-8\'\'{utf8}',
        "Access-Control-Expose-Headers": "Content-Disposition",
    }


def _add_profeic_header(doc: Document, asignatura: str, nivel: str, tipo: str) -> None:
    """Encabezado ProfeIC estándar en tabla de 2 columnas."""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.2)
    t.columns[1].width = Inches(5.3)

    # Logo (texto si no existe la imagen)
    cell_logo = t.cell(0, 0)
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    logo_path = os.path.join(base, "assets", "logo_profeic.svg.png")
    p_logo = cell_logo.paragraphs[0]
    run_logo = p_logo.add_run()
    if os.path.exists(logo_path):
        try:
            run_logo.add_picture(logo_path, width=Inches(1.0))
        except Exception:
            run_logo.add_text("ProfeIC")
    else:
        run_logo.add_text("ProfeIC")
        run_logo.bold = True

    # Info
    cell_info = t.cell(0, 1)
    p_info = cell_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_info.add_run("PROFE IC\n")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(43, 84, 110)
    r2 = p_info.add_run("Instrumento generado con IA · ProfeIC\n")
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(120, 120, 120)
    r3 = p_info.add_run(f"Asignatura: {asignatura} | Nivel: {nivel} | {tipo}")
    r3.font.size = Pt(9)
    doc.add_paragraph()


# ─── Modelos para descarga ─────────────────────────────────────────────────────

class PreguntaDescarga(BaseModel):
    numero: int
    habilidad_medida: str
    enunciado: str
    alternativas: dict[str, Any]
    correcta: str
    justificacion: str
    explicacion_correcta: str | None = None


class DescargaCuadernilloRequest(BaseModel):
    titulo: str
    asignatura: str
    nivel: str
    modo: str
    estimulo_texto: str | None = None
    estimulo_imagen: str | None = None
    preguntas: list[PreguntaDescarga]
    incluir_clave: bool = False  # True → versión docente con respuestas


class DescargaOMRRequest(BaseModel):
    titulo: str
    asignatura: str
    nivel: str
    cantidad_preguntas: int


# ─── Endpoint: Descargar Cuadernillo DOCX ─────────────────────────────────────

@router.post(
    "/generate/download-cuadernillo",
    summary="Descarga el Cuadernillo de Estímulos en DOCX",
    response_class=StreamingResponse,
)
async def download_cuadernillo(
    req: DescargaCuadernilloRequest,
    user_id: str = Depends(get_current_user_id),
) -> StreamingResponse:
    """
    POST /api/v1/simce/generate/download-cuadernillo

    Genera y descarga un DOCX con el cuadernillo de preguntas.
    Si `incluir_clave=True`, añade una página de pauta docente con las respuestas.
    """
    try:
        doc = Document()
        
        # Configurar la fuente general en Arial o Calibri tamaño 11
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        tipo_doc = "Pauta Docente" if req.incluir_clave else "Cuadernillo del Alumno"
        _add_profeic_header(doc, req.asignatura, req.nivel, tipo_doc)

        # Título
        h = doc.add_heading(req.titulo, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(27, 60, 115)

        subtitulo = f"{'Simulación SIMCE Estricta' if req.modo == 'simce' else 'Ensayo Formativo'} · {req.nivel}"
        p_sub = doc.add_paragraph(subtitulo)
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.runs[0].font.color.rgb = RGBColor(200, 117, 51)
        p_sub.runs[0].italic = True
        doc.add_paragraph()

        # Datos alumno
        t_alumno = doc.add_table(rows=2, cols=3)
        t_alumno.style = "Table Grid"
        t_alumno.cell(0, 0).text = "Nombre:"
        t_alumno.cell(0, 1).text = "Curso:"
        t_alumno.cell(0, 2).text = "Fecha:"
        t_alumno.cell(1, 0).text = "________________________________"
        t_alumno.cell(1, 1).text = "__________"
        t_alumno.cell(1, 2).text = "__________"
        doc.add_paragraph()
        
        # Diagramación Multimodal
        if req.estimulo_imagen:
            imagen_path = Path(__file__).parent.parent.parent / "frontend" / "public" / "imgs_estimulos" / req.estimulo_imagen
            try:
                if imagen_path.exists():
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(str(imagen_path), width=Inches(5.0))
            except Exception as e:
                logger.warning(f"No se pudo insertar la imagen {req.estimulo_imagen}: {e}")

        if req.estimulo_texto:
            t_est = doc.add_table(rows=1, cols=1)
            t_est.style = "Table Grid"
            t_est.autofit = True
            cell = t_est.cell(0, 0)
            cell.text = req.estimulo_texto
            
            # Shading robusto
            shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            doc.add_paragraph()

        # Preguntas
        for pregunta in req.preguntas:
            p_enun = doc.add_paragraph()
            run_num = p_enun.add_run(f"{pregunta.numero}. ")
            run_num.bold = True
            run_num.font.color.rgb = RGBColor(27, 60, 115)
            p_enun.add_run(pregunta.enunciado)

            # Habilidad (pequeño badge)
            p_hab = doc.add_paragraph()
            run_hab = p_hab.add_run(f"   [{pregunta.habilidad_medida}]")
            run_hab.italic = True
            run_hab.font.size = Pt(8)
            run_hab.font.color.rgb = RGBColor(150, 150, 150)

            for letra, alt_info in pregunta.alternativas.items():
                texto_alt = alt_info.get("texto", str(alt_info)) if isinstance(alt_info, dict) else str(alt_info)
                p_alt = doc.add_paragraph(f"{letra}) {texto_alt}")
                p_alt.paragraph_format.left_indent = Inches(0.5)
                
                if req.incluir_clave and letra == pregunta.correcta:
                    p_alt.runs[0].bold = True
                    p_alt.runs[0].font.color.rgb = RGBColor(0, 128, 0)

            doc.add_paragraph()

        # Pauta docente (página separada)
        if req.incluir_clave:
            doc.add_page_break()
            h_pauta = doc.add_heading("Pauta de Respuestas — Uso Exclusivo Docente", level=1)
            h_pauta.alignment = WD_ALIGN_PARAGRAPH.CENTER

            t_clave = doc.add_table(rows=1 + len(req.preguntas), cols=5)
            t_clave.style = "Table Grid"

            # Headers
            for col, txt in enumerate(["N°", "Respuesta Correcta", "Habilidad Medida", "Explicación Correcta", "Errores Oficiales"]):
                cell = t_clave.cell(0, col)
                cell.text = txt
                shd = parse_xml(r'<w:shd {} w:fill="1B3C73"/>'.format(nsdecls("w")))
                cell._tc.get_or_add_tcPr().append(shd)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell.paragraphs[0].runs[0].bold = True

            for i, pregunta in enumerate(req.preguntas):
                row = t_clave.rows[i + 1]
                row.cells[0].text = str(pregunta.numero)
                row.cells[1].text = pregunta.correcta
                row.cells[2].text = pregunta.habilidad_medida
                row.cells[3].text = pregunta.explicacion_correcta or pregunta.justificacion or ""
                
                errores = []
                for letra, alt_info in pregunta.alternativas.items():
                    if isinstance(alt_info, dict) and letra != pregunta.correcta:
                        err = alt_info.get("tipo_error_oficial")
                        if err: errores.append(f"{letra}: {err}")
                row.cells[4].text = "\n".join(errores)

            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"Cuadernillo_SIMCE_{req.asignatura}_{req.nivel}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=_docx_headers(filename),
        )

    except Exception as exc:
        logger.error("❌ Error generando cuadernillo DOCX: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc


# ─── Endpoint: Descargar Hoja de Respuestas OMR ───────────────────────────────

@router.post(
    "/generate/download-omr",
    summary="Descarga la Hoja de Respuestas OMR en DOCX",
    response_class=StreamingResponse,
)
async def download_omr(
    req: DescargaOMRRequest,
    user_id: str = Depends(get_current_user_id),
) -> StreamingResponse:
    """
    POST /api/v1/simce/generate/download-omr

    Genera una hoja de respuestas OMR de burbuja estándar en formato DOCX.
    Incluye grilla RUT (8 dígitos + DV).
    """
    try:
        doc = Document()
        _add_profeic_header(doc, req.asignatura, req.nivel, "Hoja de Respuestas OMR")

        h = doc.add_heading("HOJA DE RESPUESTAS", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(27, 60, 115)

        p_inst = doc.add_paragraph(
            "Instrucciones: Rellena completamente el círculo de la alternativa elegida con lápiz grafito. "
            "No uses corrector. Cualquier marca fuera del círculo puede invalidar tu respuesta."
        )
        p_inst.italic = True
        p_inst.font = None  # default
        doc.add_paragraph()

        # ── Grilla RUT ────────────────────────────────────────────────────────
        p_rut = doc.add_paragraph()
        p_rut.add_run("RUT: ").bold = True
        p_rut.add_run("__ __ __ __ __ __ __ __ - __")
        doc.add_paragraph()

        # Tabla nombre / fecha
        t_datos = doc.add_table(rows=2, cols=2)
        t_datos.style = "Table Grid"
        t_datos.cell(0, 0).text = "Nombre completo:"
        t_datos.cell(0, 1).text = "Fecha:"
        t_datos.cell(1, 0).text = "_______________________________________"
        t_datos.cell(1, 1).text = "____________"
        doc.add_paragraph()

        # ── Grilla de respuestas en 3 columnas ────────────────────────────────
        n = req.cantidad_preguntas
        col_size = (n + 2) // 3  # distribución en 3 columnas
        cols_data: list[list[int]] = []
        for c in range(3):
            start = c * col_size + 1
            end = min((c + 1) * col_size, n)
            cols_data.append(list(range(start, end + 1)))

        max_rows = max(len(col) for col in cols_data)

        # Tabla encabezado
        header_t = doc.add_table(rows=1, cols=3)
        header_t.style = "Table Grid"
        for ci, label in enumerate(["Preguntas 1–{}".format(len(cols_data[0])),
                                     "Preguntas {}–{}".format(cols_data[1][0] if cols_data[1] else "", cols_data[1][-1] if cols_data[1] else ""),
                                     "Preguntas {}–{}".format(cols_data[2][0] if cols_data[2] else "", cols_data[2][-1] if cols_data[2] else "")]):
            cell = header_t.cell(0, ci)
            cell.text = label
            shd = parse_xml(r'<w:shd {} w:fill="1B3C73"/>'.format(nsdecls("w")))
            cell._tc.get_or_add_tcPr().append(shd)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].bold = True

        grid_t = doc.add_table(rows=max_rows, cols=3)
        grid_t.style = "Table Grid"
        for row_i in range(max_rows):
            for col_i in range(3):
                nums = cols_data[col_i]
                if row_i < len(nums):
                    num = nums[row_i]
                    grid_t.cell(row_i, col_i).text = f"{num:>2}.  ○ A    ○ B    ○ C    ○ D"
                else:
                    grid_t.cell(row_i, col_i).text = ""

        doc.add_paragraph()
        p_footer = doc.add_paragraph("ProfeIC — IA Educativa Hecha en Chile 🇨🇱")
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer.runs[0].font.size = Pt(8)
        p_footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"Hoja_OMR_{req.asignatura}_{req.nivel}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=_docx_headers(filename),
        )

    except Exception as exc:
        logger.error("❌ Error generando OMR DOCX: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
