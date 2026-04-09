import google.generativeai as genai
from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
import json
import io
import os
import pandas as pd
from pypdf import PdfReader
import docx
from app.core.config import settings

# Configuración de IA para este router
if settings.GOOGLE_API_KEY:
    genai.configure(api_key=settings.GOOGLE_API_KEY)

router = APIRouter(prefix="/api/v1/mejora-continua", tags=["Mejora Continua"])

class CopilotoRequest(BaseModel):
    desafio: str
    school_id: str | None = None

@router.post("/copiloto")
async def estructurar_problema_ia(req: CopilotoRequest):
    if not req.desafio.strip():
        raise HTTPException(status_code=400, detail="El desafío no puede estar vacío.")

    try:
        model = genai.GenerativeModel(
            'gemini-2.5-flash',
            generation_config={"response_mime_type": "application/json"}
        )
        
        # Contexto Institucional
        contexto_colegio = ""
        if req.school_id:
            try:
                from app.db.supabase import supabase as db_supabase
                school_resp = db_supabase.table("schools").select("*").eq("id", req.school_id).maybe_single().execute()
                if school_resp.data:
                    s = school_resp.data
                    contexto_colegio = f"""
DATOS DEL COLEGIO:
- Nombre: {s.get('name')}
- Vulnerabilidad (Prioritarios): {s.get('priority_pct')}%
- Alumnos PIE: {s.get('pie_neet_count', 0) + s.get('pie_neep_count', 0)} alumnos.
- Asistencia Promedio: {s.get('attendance_avg')}%
- Nivel Socioeconómico: {s.get('socioeconomic_level')}
"""
            except Exception as e_ctx:
                print(f"⚠️ Error cargando contexto colegio en Copiloto: {e_ctx}")

        system_instruction = f"""Eres un 'Consultor Estratégico en Gestión Escolar'. Tu objetivo es estructurar el desafío ingresado en un Plan de Mejora Continua. Debes generar exactamente 3 fases lógicas.
{contexto_colegio}
Debes retornar estrictamente un JSON con esta estructura (nada más):
{{
  "title": "Título formal de la meta estratégica",
  "description": "Breve descripción",
  "phases": [
    {{
      "title": "Nombre de la fase",
      "indicators": [
        {{
          "description": "Indicador SMART", 
          "target_value": 85, 
          "metric_type": "percentage", 
          "data_source": "radar_360"
        }}
      ],
      "enablers": [
        {{
          "description": "Recurso necesario", 
          "resource_type": "training"
        }}
      ]
    }}
  ]
}}
Reglas estrictas para los valores de ciertos campos:
- `metric_type` DEBE SER SOLAMENTE 'percentage', 'average_score', o 'count'.
- `data_source` DEBE SER SOLAMENTE 'manual' o 'radar_360'.
- `resource_type` DEBE SER SOLAMENTE 'time', 'training', 'technology', o 'material'.
Asegúrate de respetar fielmente esta estructura de JSON, ya que se usará para popular la UI directamente."""

        prompt = f"{system_instruction}\n\nDesafío del usuario: {req.desafio}"
        
        # Como no se importó el aio module, para simplificar e ir a lo seguro
        # usamos generate_content en vez de la versión asíncrona si no estamos seguros de si la versión SDK tiene aio nativamente fácil aquí.
        # Las llamadas a generate_content_async son con model.generate_content_async()
        response = await model.generate_content_async(prompt)
        
        if not response.text:
            raise HTTPException(status_code=500, detail="El modelo no devolvió contenido.")

        result_json = json.loads(response.text)
        return result_json

    except json.JSONDecodeError as e:
        print(f"[ERROR] Error decodificando JSON de Gemini: {e}")
        raise HTTPException(status_code=500, detail="El formato devuelto por la IA fue inválido.")
    except Exception as e:
        print(f"[ERROR] Error Inesperado en copiloto de PME: {e}")
        raise HTTPException(status_code=500, detail=f"Ocurrió un error inesperado trabajando con IA: {str(e)}")


async def extract_text_from_file(file: UploadFile) -> str:
    """Extrae texto base de diferentes formatos de archivo de forma robusta."""
    filename = file.filename.lower()
    content = await file.read()
    text = ""
    
    try:
        if filename.endswith(".pdf"):
            pdf = PdfReader(io.BytesIO(content))
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
                
        elif filename.endswith(".xlsx") or filename.endswith(".csv"):
            if filename.endswith(".csv"):
                df = pd.read_csv(io.BytesIO(content))
            else:
                df = pd.read_excel(io.BytesIO(content))
            text = df.to_string()
            
        else:
            raise ValueError("Formato de archivo no soportado. Usa PDF, DOCX, XLSX o CSV.")
            
        return text.strip()
    except Exception as e:
        raise ValueError(f"No se pudo extraer texto del archivo: {str(e)}")


@router.post("/importar-documento")
async def importar_documento_pme(file: UploadFile = File(...)):
    if not file:
        raise HTTPException(status_code=400, detail="No se ha proporcionado ningún archivo.")
        
    try:
        raw_text = await extract_text_from_file(file)
        
        if len(raw_text) < 50:
            raise HTTPException(status_code=400, detail="El documento parece estar vacío o su contenido no pudo ser leído de forma correcta.")

        # Limitamos el texto para no exceder tokens (100k chars suele ser seguro para Gemini 1.5/2.5 Flash, pero es mejor ser cautos)
        raw_text = raw_text[:80000]

        model = genai.GenerativeModel(
            'gemini-2.5-flash',
            generation_config={"response_mime_type": "application/json"}
        )
        
        system_instruction = """Eres un analizador experto de datos escolares y consultor estratégico. 
A continuación recibirás el texto crudo extraído de un documento oficial de un Plan de Mejora Educativa (PME).
Tu trabajo es interpretar su contenido y estructurarlo ESTRICTAMENTE en nuestro formato JSON.

### DIMENSIONES PME (CHILE)
Debes clasificar cada fase o acción en una de estas 4 dimensiones obligatorias:
1. "Gestión Pedagógica"
2. "Liderazgo"
3. "Convivencia Escolar"
4. "Gestión de Recursos"

### ESTRUCTURA JSON REQUERIDA:
{
  "title": "Meta estratégica global",
  "description": "Breve resumen estratégico",
  "phases": [
    {
      "title": "[DIMENSIÓN] Nombre descriptivo de la acción/fase",
      "dimension": "Una de las 4 dimensiones anteriores",
      "indicators": [
        {
          "description": "Indicador SMART", 
          "target_value": 100, 
          "metric_type": "percentage", 
          "data_source": "manual"
        }
      ],
      "enablers": [
        {
          "description": "Recurso o medio", 
          "resource_type": "material",
          "estimated_cost": 500000 // Infiere el costo si se menciona, si no, usa 0
        }
      ]
    }
  ]
}

### REGLAS CRÍTICAS:
- El campo `title` de la FASE debe comenzar con la dimensión entre corchetes, ej: "[Gestión Pedagógica] Taller de DUA".
- `metric_type` DEBE SER SOLAMENTE 'percentage', 'average_score', o 'count'.
- `data_source` DEBE SER SOLAMENTE 'manual' o 'radar_360'.
- `resource_type` DEBE SER SOLAMENTE 'time', 'training', 'technology', o 'material'.
- Si el documento menciona montos de inversión, cífrelos en `estimated_cost` como números enteros.

Genera hasta 4 fases bien estructuradas basadas únicamente en el contenido del texto."""

        prompt = f"{system_instruction}\n\n=== TEXTO DEL PME ===\n{raw_text}\n=====================\n"
        
        response = await model.generate_content_async(prompt)
        
        if not response.text:
            raise HTTPException(status_code=500, detail="El modelo IA no devolvió contenido.")

        result_json = json.loads(response.text)
        return result_json

    except ValueError as ve:
        print(f"[ERROR] Importando documento PME: {ve}")
        raise HTTPException(status_code=400, detail=str(ve))
    except json.JSONDecodeError as e:
        print(f"[ERROR] JSON IA malformado: {e}")
        raise HTTPException(status_code=500, detail="La IA generó una estructura inválida.")
    except Exception as e:
        print(f"[ERROR] Inesperado importar-documento: {e}")
        raise HTTPException(status_code=500, detail=f"Fallo en el servidor al procesar tu documento: {str(e)}")

from typing import Any

@router.post("/generar-informe")
async def generar_informe_ia(data: Any = None):
    try:
        # Forzamos configuración en cada llamada por si el contexto global se pierde
        if settings.GOOGLE_API_KEY:
            genai.configure(api_key=settings.GOOGLE_API_KEY)
            
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            generation_config={"response_mime_type": "application/json"}
        )
        
        sys_prompt = """Eres un Consultor Estratégico de Alto Nivel evaluando un Plan de Mejora Educativa. Analiza los datos de las metas, fases e indicadores provistos. NO escribas párrafos largos. Retorna tu análisis ESTRICTAMENTE en el siguiente formato JSON:
{
  "health_status": "Crítico" | "Estable" | "Óptimo",
  "health_score": <número del 0 al 100 evaluando el avance global>,
  "executive_summary": "Un solo párrafo de máximo 3 líneas resumiendo la situación actual de manera directiva y al grano.",
  "critical_alerts": [
    {"title": "Título", "description": "Descripción"}
  ],
  "strategic_recommendations": [
    {"action": "Verbo", "detail": "Detalle"}
  ],
  "positive_highlights": [
    "Viñeta"
  ],
  "pending_tasks": [
    {"task": "Tarea", "responsible": "Responsable"}
  ]
}
Sé crudo, analítico y preciso."""
        
        payload_str = json.dumps(data, indent=2, ensure_ascii=False)
        prompt = f"{sys_prompt}\n\n[DATOS DEL PME PARA ANALIZAR]:\n{payload_str}"
        
        response = await model.generate_content_async(prompt)
        
        # Validación de respuesta segura
        if not response.text:
            print(f"⚠️ [IA PME] Respuesta vacía. Candidatos: {response.candidates}")
            raise HTTPException(status_code=500, detail="La IA no devolvió texto. Posible bloqueo de seguridad o cuota.")

        return json.loads(response.text)

    except Exception as e:
        import traceback
        error_msg = f"Error en Informe IA: {str(e)}"
        print(f"❌ {error_msg}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=error_msg)


from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_cell_background(cell, fill, color=None, val=None):
    """
    Establece el fondo de una celda en python-docx.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def _add_page_number(run):
    """
    Añade numeración de página dinámica.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

class ComparativaRequest(BaseModel):
    schools_data: list # Lista de colegios con sus metas y fases

@router.post("/comparativa-holding")
async def generar_comparativa_holding(req: ComparativaRequest):
    try:
        if not req.schools_data:
            raise HTTPException(status_code=400, detail="No hay datos de colegios para comparar.")

        model = genai.GenerativeModel(
            'gemini-2.5-flash',
            generation_config={"response_mime_type": "application/json"}
        )
        
        system_instruction = """Eres un Analista Estratégico de Redes Educacionales (Holding).
Tu misión es recibir un conjunto de Planes de Mejora Educativa (PME) de diferentes colegios y generar una visión comparativa de alto nivel.
Debes identificar sinergias, brechas críticas y oportunidades de optimización de recursos.

Retorna ESTRICTAMENTE un JSON con esta estructura:
{
  "resumen_ejecutivo": "Un párrafo sobre el estado general de la red.",
  "comunalidades": [
    {"tema": "Título", "descripcion": "Por qué es común y cuántos colegios lo comparten"}
  ],
  "diferencias_clave": [
    {"colegio": "Nombre", "aspecto_unico": "Qué está haciendo diferente al resto"}
  ],
  "analisis_presupuestario": {
    "total_red": 12000000,
    "distribucion_sugerida": "Breve nota sobre si el gasto está bien enfocado"
  },
  "oportunidades_sinergia": [
    "Ej: El Colegio A y B pueden compartir capacitaciones de DUA porque ambos tienen metas similares."
  ]
}"""

        payload_str = json.dumps(req.schools_data, indent=2, ensure_ascii=False)
        prompt = f"{system_instruction}\n\n[DATOS DE LOS COLEGIOS A COMPARAR]:\n{payload_str}"
        
        response = await model.generate_content_async(prompt)
        
        if not response.text:
            raise HTTPException(status_code=500, detail="La IA no pudo procesar la comparativa.")

        return json.loads(response.text)

    except Exception as e:
        print(f"[ERROR] Comparativa Holding: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/descargar-word")
async def descargar_informe_word(payload: dict):
    try:
        informe = payload.get("informe", {})
        dashboard_data = payload.get("dashboard_data", [])
        
        doc = Document()
        
        # --- PORTADA ---
        # Logo Centrado
        logo_path = "/Users/renealvarezpinones/Downloads/PROFEIC_GITHUB_01_2/backend/assets/logo_profeic..png"
        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_logo = p_logo.add_run()
            r_logo.add_picture(logo_path, width=Inches(2.5))
        
        doc.add_paragraph("\n\n\n")
        
        # Título Portada
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run("INFORME DIRECTIVO ESTRATÉGICO\nPLAN DE MEJORA EDUCATIVA (PME)")
        r_title.font.bold = True
        r_title.font.size = Pt(24)
        r_title.font.color.rgb = RGBColor(0x02, 0x3e, 0x8a)
        
        doc.add_paragraph("\n\n")
        
        # Subtítulo Portada
        p_subtitle = doc.add_paragraph()
        p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_subtitle = p_subtitle.add_run("Auditoría de Gestión Académica y Copiloto IA")
        r_subtitle.font.size = Pt(14)
        r_subtitle.font.italic = True
        r_subtitle.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        
        doc.add_paragraph("\n\n\n\n")
        
        # Metadata Portada
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_meta = p_meta.add_run(f"Generado el {pd.Timestamp.now().strftime('%d de %B, %Y')}\nProfeIC Analytics Service")
        r_meta.font.size = Pt(10)
        
        doc.add_page_break()

        # --- PIE DE PÁGINA (Paginación) ---
        footer = doc.sections[0].footer
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = p_footer.add_run("Página ")
        _add_page_number(run_footer)

        # --- SECCIÓN 1: SALUD DEL PLAN ---
        doc.add_heading("I. ESTADO GENERAL DE SALUD", level=1)
        doc.add_paragraph("Resumen cuantitativo del avance global del Plan de Mejora Educativa basado en indicadores críticos.")
        
        if informe:
            stable = doc.add_table(rows=1, cols=2)
            stable.width = Inches(6)
            cells = stable.rows[0].cells
            
            # Celda Score con Fondo
            _set_cell_background(cells[0], "F1F5F9")
            p0 = cells[0].add_paragraph()
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(f"SCORE GLOBAL\n{informe.get('health_score')}%")
            r0.font.bold = True
            r0.font.size = Pt(20)
            r0.font.color.rgb = RGBColor(0x02, 0x3e, 0x8a)
            
            # Celda Estado
            status = informe.get('health_status', 'N/A')
            _set_cell_background(cells[1], "F1F5F9")
            p1 = cells[1].add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(f"CALIFICACIÓN\n{status.upper()}")
            r1.font.bold = True
            r1.font.size = Pt(18)
            if status == "Óptimo": r1.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
            elif status == "Estable": r1.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)
            else: r1.font.color.rgb = RGBColor(0xef, 0x44, 0x44)

        doc.add_paragraph("\n")
        
        # --- SECCIÓN 2: ANÁLISIS ESTRATÉGICO ---
        if informe:
            doc.add_heading("II. ANÁLISIS ESTRATÉGICO (CO-PILOT)", level=1)
            p_exec = doc.add_paragraph(informe.get('executive_summary', ''))
            p_exec.style = 'Intense Quote'
            
            if informe.get("pending_tasks"):
                doc.add_heading("MONITOREO DE PENDIENTES CRÍTICOS", level=2)
                doc.add_paragraph("Fases o tareas que presentan estancamiento y requieren atención inmediata.")
                
                ttable = doc.add_table(rows=1, cols=2)
                ttable.style = 'Table Grid'
                hdr = ttable.rows[0].cells
                hdr[0].text = 'Acción / Tarea Pendiente'
                hdr[1].text = 'Responsable Directo'
                for h in hdr: 
                    h.paragraphs[0].runs[0].font.bold = True
                    _set_cell_background(h, "E2E8F0")
                
                for t in informe.get("pending_tasks", []):
                    row = ttable.add_row().cells
                    row[0].text = str(t.get("task", ""))
                    row[1].text = str(t.get("responsible", ""))
            
            doc.add_paragraph("\n")

        # --- SECCIÓN 3: DETALLE DE IMPLEMENTACIÓN ---
        doc.add_page_break()
        doc.add_heading("III. DETALLE DE IMPLEMENTACIÓN (ACCIONES PME)", level=1)
        doc.add_paragraph("Desglose de metas estratégicas, fases de implementación y monitoreo de indicadores.")
        
        if not dashboard_data:
            doc.add_paragraph("No se detectaron datos de ejecución en el plan actual.")
        else:
            for i, goal in enumerate(dashboard_data, 1):
                doc.add_heading(f"Meta Estratégica {i}: {goal.get('title', 'Sin título')}", level=2)
                
                gtable = doc.add_table(rows=1, cols=3)
                gtable.style = 'Table Grid'
                hdr = gtable.rows[0].cells
                hdr[0].text = 'Fase de Implementación'
                hdr[1].text = 'Responsable'
                hdr[2].text = 'Progreso'
                for h in hdr: 
                    h.paragraphs[0].runs[0].font.bold = True
                    _set_cell_background(h, "F8FAFC")
                
                for phase in goal.get("implementation_phases", []):
                    row = gtable.add_row().cells
                    row[0].text = f"{phase.get('title')} ({phase.get('status')})"
                    row[1].text = str(phase.get('leader_name', 'No asignado'))
                    row[2].text = f"{phase.get('average_score', 0)}%"
                    
                doc.add_paragraph("\n")

        # --- SECCIÓN 4: RECOMENDACIONES ---
        if informe and informe.get("strategic_recommendations"):
            doc.add_heading("IV. RECOMENDACIONES DE OPTIMIZACIÓN", level=1)
            for rec in informe.get("strategic_recommendations", []):
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(f"{rec.get('action')}: ")
                r.font.bold = True
                p.add_run(rec.get('detail'))

        # --- CIERRE Y FIRMAS ---
        doc.add_page_break()
        doc.add_heading("VALIDACIÓN INSTITUCIONAL", level=1)
        doc.add_paragraph("\n\n\n\n")
        
        ftable = doc.add_table(rows=1, cols=2)
        cells = ftable.rows[0].cells
        
        p1 = cells[0].add_paragraph("__________________________\nFirma Director(a)\nFecha: ___/___/___")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p2 = cells[1].add_paragraph("__________________________\nFirma Coordinador(a) PME\nFecha: ___/___/___")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        return StreamingResponse(
            buf, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Informe_Directivo_PME_ProfeIC.docx"}
        )

    except Exception as e:
        import traceback
        print(f"❌ Error crítico en Word Export: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Fallo en generación de reporte: {str(e)}")

async def _descargar_word_legacy(informe_md: str):
    doc = Document()
    doc.add_heading("Informe Directivo PME", 0)
    doc.add_paragraph(informe_md)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ─── Gantt PDF Export (reportlab) ────────────────────────────────────────────

from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable, Image
)
from reportlab.platypus import KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from fastapi.responses import StreamingResponse
from pydantic import BaseModel as PydanticBaseModel
from typing import List

class GanttRow(PydanticBaseModel):
    objetivo: str
    fase: str
    lider: str
    estado: str
    inicio: str
    fin: str
    avance: int

class GanttExportRequest(PydanticBaseModel):
    rows: List[GanttRow]
    filter_goal_label: str = "Todos los objetivos"
    filter_leader_label: str = "Todos los líderes"


STATUS_COLORS_RL = {
    "Pendiente":    colors.HexColor("#64748B"),
    "En Ejecución": colors.HexColor("#3B82F6"),
    "Logrado":      colors.HexColor("#10B981"),
}


@router.post("/exportar-gantt-pdf")
async def exportar_gantt_pdf(req: GanttExportRequest):
    """Genera un PDF de alta calidad del Roadmap PME usando reportlab."""
    try:
        buf = io.BytesIO()
        page_w, page_h = landscape(A4)
        margin = 15 * mm

        doc = SimpleDocTemplate(
            buf,
            pagesize=landscape(A4),
            leftMargin=margin, rightMargin=margin,
            topMargin=margin, bottomMargin=margin,
        )

        styles = getSampleStyleSheet()
        SLATE = colors.HexColor("#0F172A")
        SLATE_400 = colors.HexColor("#94A3B8")
        SLATE_100 = colors.HexColor("#F1F5F9")
        BLUE_400 = colors.HexColor("#60A5FA")

        # Cell paragraph styles
        cell_style = ParagraphStyle("cell", fontName="Helvetica", fontSize=8, leading=10, textColor=colors.HexColor("#1E293B"))
        sub_style  = ParagraphStyle("sub",  fontName="Helvetica", fontSize=6.5, leading=8, textColor=colors.HexColor("#64748B"))
        hdr_style  = ParagraphStyle("hdr",  fontName="Helvetica-Bold", fontSize=7, leading=9, textColor=colors.white, alignment=TA_CENTER)
        status_style = ParagraphStyle("st", fontName="Helvetica-Bold", fontSize=6.5, leading=8, textColor=colors.white, alignment=TA_CENTER)

        elements = []

        # ── Header band (simulated with a single-cell table) ─────────────────────
        import os
        logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "logo_profeic..png")
        logo_flowable = ""
        if os.path.exists(logo_path):
            try:
                logo_flowable = Image(logo_path, width=15*mm, height=15*mm)
            except Exception:
                pass

        header_data = [[
            logo_flowable if logo_flowable else "",
            Paragraph(f'<font size="16"><b>ROADMAP ESTRATÉGICO PME 2026</b></font>', ParagraphStyle("h", fontName="Helvetica-Bold", fontSize=16, textColor=colors.white)),
            Paragraph(f'<font size="7">Filtro: {req.filter_goal_label}  •  Líder: {req.filter_leader_label}<br/>{len(req.rows)} fases activas</font>',
                      ParagraphStyle("sub_h", fontName="Helvetica", fontSize=7, textColor=SLATE_400, alignment=TA_RIGHT))
        ]]
        header_tbl = Table(header_data, colWidths=[20 * mm, page_w * 0.55 - margin - 20 * mm, page_w * 0.45 - margin])
        header_tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, -1), SLATE),
            ("TOPPADDING",   (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 8),
            ("LEFTPADDING",  (0, 0), (0, -1),  10),
            ("RIGHTPADDING", (-1, 0), (-1, -1), 10),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elements.append(header_tbl)
        elements.append(Spacer(1, 6 * mm))

        # ── Column headers row ────────────────────────────────────────────────────
        usable_w = page_w - 2 * margin
        col_widths = [
            usable_w * 0.28,   # Objetivo
            usable_w * 0.23,   # Fase
            usable_w * 0.16,   # Líder
            usable_w * 0.10,   # Estado
            usable_w * 0.08,   # Inicio
            usable_w * 0.08,   # Fin
            usable_w * 0.07,   # Avance
        ]
        headers = ["OBJETIVO", "FASE / ACCIÓN", "RESPONSABLE", "ESTADO", "INICIO", "FIN", "AVANCE"]
        header_row = [Paragraph(h, hdr_style) for h in headers]

        # ── Data rows ─────────────────────────────────────────────────────────────
        table_data = [header_row]
        for i, row in enumerate(req.rows):
            status_color = STATUS_COLORS_RL.get(row.estado, SLATE_400)
            status_cell = Paragraph(row.estado, status_style)

            # Progress bar as a mini table
            filled = max(int(row.avance * col_widths[6] / 100) - 2, 0)
            empty  = max(int(col_widths[6]) - filled - 2, 0)
            bar_data = [[""]]
            bar = Table(bar_data, colWidths=[filled + empty])
            bar_inner_filled = Table([[""]], colWidths=[filled])
            bar_inner_filled.setStyle(TableStyle([
                ("BACKGROUND", (0,0),(-1,-1), colors.HexColor("#3B82F6")),
                ("TOPPADDING", (0,0),(-1,-1), 3),
                ("BOTTOMPADDING", (0,0),(-1,-1), 3),
            ]))
            avance_text = Paragraph(f"<b>{row.avance}%</b>", ParagraphStyle("av", fontName="Helvetica-Bold", fontSize=7, textColor=colors.HexColor("#1E293B"), alignment=TA_CENTER))

            table_data.append([
                Paragraph(str(row.objetivo)[:120], cell_style),
                Paragraph(str(row.fase)[:100],     cell_style),
                Paragraph(str(row.lider)[:40],    sub_style),
                Paragraph(f"<b>{row.estado.upper()}</b>", status_style),
                Paragraph(str(row.inicio),        sub_style),
                Paragraph(str(row.fin),           sub_style),
                avance_text,
            ])

        # ── Build main table ──────────────────────────────────────────────────────
        main_table = Table(table_data, colWidths=col_widths, repeatRows=1)

        row_styles = [
            ("BACKGROUND",    (0, 0), (-1, 0),    SLATE),
            ("TEXTCOLOR",     (0, 0), (-1, 0),    colors.white),
            ("FONTNAME",      (0, 0), (-1, 0),    "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, 0),    7),
            ("TOPPADDING",    (0, 0), (-1, 0),    6),
            ("BOTTOMPADDING", (0, 0), (-1, 0),    6),
            ("GRID",          (0, 0), (-1, -1),   0.3, colors.HexColor("#E2E8F0")),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1),   [colors.white, colors.HexColor("#F8FAFC")]),
            ("VALIGN",        (0, 0), (-1, -1),   "MIDDLE"),
            ("TOPPADDING",    (0, 1), (-1, -1),   5),
            ("BOTTOMPADDING", (0, 1), (-1, -1),   5),
            ("LEFTPADDING",   (0, 0), (-1, -1),   4),
            ("RIGHTPADDING",  (0, 0), (-1, -1),   4),
        ]

        for i, row in enumerate(req.rows, start=1):
            sc = STATUS_COLORS_RL.get(row.estado, SLATE_400)
            row_styles.append(("BACKGROUND", (3, i), (3, i), sc))
            row_styles.append(("TEXTCOLOR",  (3, i), (3, i), colors.white))
            row_styles.append(("ALIGN",      (3, i), (3, i), "CENTER"))

        main_table.setStyle(TableStyle(row_styles))
        elements.append(main_table)

        elements.append(Spacer(1, 8 * mm))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=SLATE_100))
        elements.append(Spacer(1, 3 * mm))

        # ── Footer ────────────────────────────────────────────────────────────────
        from datetime import datetime
        fecha = datetime.now().strftime("%-d de %B de %Y").capitalize()
        footer = Table([[
            Paragraph(f'<font size="7" color="#94A3B8">ProfeIC 360  •  Mejorando Juntos  •  {fecha}</font>',
                      ParagraphStyle("ftr", fontName="Helvetica", fontSize=7, textColor=SLATE_400)),
            Paragraph('<font size="7" color="#94A3B8">Sistema de Gestión Estratégica — Mejorando Juntos</font>',
                      ParagraphStyle("ftr2", fontName="Helvetica", fontSize=7, textColor=SLATE_400, alignment=TA_RIGHT)),
        ]], colWidths=[usable_w * 0.5, usable_w * 0.5])
        footer.setStyle(TableStyle([("TOPPADDING", (0,0),(-1,-1), 0), ("BOTTOMPADDING", (0,0),(-1,-1), 0)]))
        elements.append(footer)

        doc.build(elements)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=MejorandoJuntos_{datetime.now().year}.pdf"}
        )
    except Exception as e:
        import traceback
        print(f"❌ Error crítico en PDF Pro Export: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Fallo en generación de PDF: {str(e)}")
