"""
Motor de renderizado de evaluaciones para ProfeIC.
Convierte estructuras AssessmentObject validadas en documentos Word y Excel.
"""

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from src.prompts import AssessmentObject, QuestionItem


def create_word_from_assessment(assessment_data: AssessmentObject) -> BytesIO:
    """
    Crea un documento Word (.docx) a partir de una estructura AssessmentObject validada.
    
    Args:
        assessment_data: Objeto AssessmentObject con los datos de la evaluación
    
    Returns:
        BytesIO con el contenido del archivo .docx
    """
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading(assessment_data.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Objetivo de aprendizaje
    doc.add_paragraph(f"Objetivo de Aprendizaje: {assessment_data.learning_objective}")
    doc.add_paragraph()  # Espacio
    
    # Agregar cada pregunta
    for i, question in enumerate(assessment_data.questions, 1):
        # Número de pregunta
        question_para = doc.add_paragraph()
        question_para.add_run(f"{i}. ").bold = True
        question_para.add_run(question.question_text)
        
        # Información de calibración (en cursiva, más pequeño)
        calibration_info = doc.add_paragraph(
            f"Bloom: {question.bloom_level} | DOK: {question.dok_level}",
            style='Intense Quote'
        )
        calibration_info.runs[0].font.size = Pt(9)
        calibration_info.runs[0].italic = True
        
        # Opciones de respuesta
        for j, option in enumerate(question.options, 1):
            option_para = doc.add_paragraph(f"   {chr(96 + j)}. {option}", style='List Bullet')
        
        # Respuesta correcta (en negrita)
        correct_para = doc.add_paragraph()
        correct_para.add_run("Respuesta correcta: ").bold = True
        correct_para.add_run(question.correct_answer)
        
        # Justificación de calibración
        justification_para = doc.add_paragraph()
        justification_para.add_run("Justificación: ").bold = True
        justification_para.add_run(question.justification_for_calibration)
        justification_para.runs[1].font.size = Pt(9)
        justification_para.runs[1].italic = True
        
        doc.add_paragraph()  # Espacio entre preguntas
    
    # Guardar en BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer


def create_excel_from_assessment(assessment_data: AssessmentObject) -> BytesIO:
    """
    Crea un archivo Excel (.xlsx) a partir de una estructura AssessmentObject validada.
    
    Args:
        assessment_data: Objeto AssessmentObject con los datos de la evaluación
    
    Returns:
        BytesIO con el contenido del archivo .xlsx
    """
    # Crear DataFrame con las preguntas
    rows = []
    for i, question in enumerate(assessment_data.questions, 1):
        row = {
            "N°": i,
            "Pregunta": question.question_text,
            "Bloom": question.bloom_level,
            "DOK": question.dok_level,
            "Opción A": question.options[0] if len(question.options) > 0 else "",
            "Opción B": question.options[1] if len(question.options) > 1 else "",
            "Opción C": question.options[2] if len(question.options) > 2 else "",
            "Opción D": question.options[3] if len(question.options) > 3 else "",
            "Opción E": question.options[4] if len(question.options) > 4 else "",
            "Respuesta Correcta": question.correct_answer,
            "Justificación": question.justification_for_calibration
        }
        rows.append(row)
    
    df = pd.DataFrame(rows)
    
    # Crear Excel en memoria
    excel_buffer = BytesIO()
    
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        # Hoja principal con preguntas
        df.to_excel(writer, sheet_name='Preguntas', index=False)
        
        # Hoja con información general
        info_df = pd.DataFrame({
            "Campo": ["Título", "Objetivo de Aprendizaje", "Total de Preguntas"],
            "Valor": [
                assessment_data.title,
                assessment_data.learning_objective,
                len(assessment_data.questions)
            ]
        })
        info_df.to_excel(writer, sheet_name='Información', index=False)
        
        # Hoja con resumen de calibración
        bloom_counts = {}
        dok_counts = {}
        for q in assessment_data.questions:
            bloom_counts[q.bloom_level] = bloom_counts.get(q.bloom_level, 0) + 1
            dok_counts[q.dok_level] = dok_counts.get(q.dok_level, 0) + 1
        
        calibration_df = pd.DataFrame({
            "Nivel Bloom": list(bloom_counts.keys()),
            "Cantidad": list(bloom_counts.values())
        })
        calibration_df.to_excel(writer, sheet_name='Resumen Bloom', index=False)
        
        dok_df = pd.DataFrame({
            "Nivel DOK": list(dok_counts.keys()),
            "Cantidad": list(dok_counts.values())
        })
        dok_df.to_excel(writer, sheet_name='Resumen DOK', index=False)
    
    excel_buffer.seek(0)
    return excel_buffer


def create_word_from_dict(assessment_dict: dict) -> BytesIO:
    """
    Crea un documento Word a partir de un diccionario (útil para datos de API).
    
    Args:
        assessment_dict: Diccionario con estructura de AssessmentObject
    
    Returns:
        BytesIO con el contenido del archivo .docx
    """
    assessment_data = AssessmentObject(**assessment_dict)
    return create_word_from_assessment(assessment_data)


def create_excel_from_dict(assessment_dict: dict) -> BytesIO:
    """
    Crea un archivo Excel a partir de un diccionario (útil para datos de API).
    
    Args:
        assessment_dict: Diccionario con estructura de AssessmentObject
    
    Returns:
        BytesIO con el contenido del archivo .xlsx
    """
    assessment_data = AssessmentObject(**assessment_dict)
    return create_excel_from_assessment(assessment_data)


def create_word_from_plan(plan_text: str, metadata: dict) -> BytesIO:
    """
    Crea un documento Word a partir de un texto de plan de clase.
    
    Args:
        plan_text: Texto del plan de clase
        metadata: Metadatos (asignatura, nivel, objetivo)
    
    Returns:
        BytesIO con el contenido del archivo .docx
    """
    doc = Document()
    
    # Título
    title = doc.add_heading('Plan de Clase', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadatos
    if metadata:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Asignatura / Nivel'
        hdr_cells[1].text = f"{metadata.get('asignatura', '')} - {metadata.get('nivel', '')}"
        
        row = table.add_row().cells
        row[0].text = 'Objetivo'
        row[1].text = metadata.get('objetivo', '')
    
    doc.add_paragraph()
    
    # Contenido del plan
    # Intentar parsear secciones si el texto tiene formato markdown
    for line in plan_text.split('\\n'):
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line.replace('- ', '').replace('* ', ''), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

