import os
import json
import random
import traceback
import asyncio
import io
import re
import unicodedata
from typing import List, Dict, Any, Optional
from pathlib import Path
from urllib.parse import quote
import glob
from fastapi import APIRouter, HTTPException, Depends, File, UploadFile
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from supabase import create_client, Client
from google import genai
from google.genai import types

# Para exportación DOCX
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

# Importar el calculador matemático
from simce_parser import calcular_distribucion
# Importar dependencia de autenticación
from routers.deps import get_current_user_id_optional

router = APIRouter(prefix="/api/v1/simce", tags=["SIMCE"])

# Configuración de Rutas y Supabase
_ASSETS_DIR = Path(__file__).resolve().parent / "assets"
_ESTIMULOS_JSON = _ASSETS_DIR / "banco_estimulos_unificado.json"
INDICADORES_PATH = _ASSETS_DIR / "indicadores_curriculum_PERFECTO.json"
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

# Mapeo de asignaturas
SUBJECT_MAPS = {
    "Lenguaje y Comunicación - Lectura": {"banco": "Lengua y Literatura", "json": "Lengua y Literatura"},
    "Historia, Geografía y Ciencias Sociales": {"banco": "Historia y Geografía", "json": "Historia, Geografía y Ciencias Sociales"},
    "Matemática": {"banco": "Matemática", "json": "Matemática"},
    "Ciencias Naturales": {"banco": "Ciencias Naturales", "json": "Ciencias Naturales"}
}

# Inicializar cliente Gemini
try:
    ai_client = genai.Client()
except Exception as e:
    print(f"Advertencia: No se pudo inicializar Gemini Client. Error: {e}")

# --- Modelos Pydantic ---

class SimceRequest(BaseModel):
    nivel: str
    asignatura: str
    cantidad_preguntas: int
    modo: str = "simce"
    oas_seleccionados: List[str] = []

class Pregunta(BaseModel):
    numero: int
    habilidad_medida: str
    enunciado: str
    alternativas: Dict[str, Any]
    correcta: str
    justificacion: str
    estimulo_texto: Optional[str] = None
    estimulo_imagen: Optional[str] = None

class BloqueEnsayo(BaseModel):
    tipo_estimulo: str
    contenido_estimulo: str
    preguntas: List[Pregunta]
    ruta_local: Optional[str] = None

class SimceResponse(BaseModel):
    titulo: str
    asignatura: str
    nivel: str
    cantidad_preguntas: int
    modo: str
    blueprint_resumen: str
    estimulo_texto: Optional[str] = None
    estimulo_imagen: Optional[str] = None
    ensayo: List[BloqueEnsayo]
    preguntas: List[Pregunta] = []

# Modelos para Descarga
class DescargaCuadernilloRequest(BaseModel):
    titulo: str
    asignatura: str
    nivel: str
    modo: str
    estimulo_texto: Optional[str] = None
    estimulo_imagen: Optional[str] = None
    preguntas: Optional[List[Pregunta]] = None
    ensayo: Optional[List[BloqueEnsayo]] = None
    incluir_clave: bool = False

class DescargaOMRRequest(BaseModel):
    titulo: str
    asignatura: str
    nivel: str
    cantidad_preguntas: int

# --- Funciones de Utilidad ---

def canonical(s: str) -> str:
    if not s: return ""
    s = s.lower().strip()
    s = s.replace("°", "").replace("º", "").replace(".", "").replace(" ", "").replace("-", "")
    s = s.replace("iv", "4").replace("iii", "3").replace("ii", "2").replace("i", "1")
    s = s.replace("lenguaje", "lengua")
    return s

def extract_json(text: str) -> Any:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    match = re.search(r"(\{.*\}|\[.*\])", text, re.DOTALL)
    if match:
        try: return json.loads(match.group(1))
        except: pass
    try: return json.loads(text)
    except: return None

def load_local_json(path: Path) -> List[Dict[str, Any]]:
    if not path.exists(): return []
    try:
        with open(path, "r", encoding="utf-8") as f: return json.load(f)
    except: return []

def get_indicadores_locales(asignatura_oficial: str, nivel: str) -> str:
    data = load_local_json(INDICADORES_PATH)
    
    # Normalización vía Map
    norm_asig = canonical(asignatura_oficial)
    if asignatura_oficial in SUBJECT_MAPS:
        norm_asig = canonical(SUBJECT_MAPS[asignatura_oficial].get("json", asignatura_oficial))
    
    norm_nivel = canonical(nivel)
        
    matches = [
        item for item in data 
        if canonical(item.get("asignatura", "")) == norm_asig 
        and canonical(item.get("nivel", "")) == norm_nivel
    ]
    
    # Intento 2: Búsqueda parcial si falla la exacta
    if not matches:
        matches = [
            item for item in data 
            if (norm_asig in canonical(item.get("asignatura", "")) or canonical(item.get("asignatura", "")) in norm_asig)
            and canonical(item.get("nivel", "")) == norm_nivel
        ]

    if not matches: return ""
    all_inds = []
    for m in matches: all_inds.extend(m.get("indicadores", []))
    return "\n- ".join(all_inds[:20])

def _get_random_stimulus(asignatura: str, nivel: str) -> dict:
    if not _ESTIMULOS_JSON.exists():
        raise HTTPException(
            status_code=500, 
            detail=f"CRÍTICO: No se encontró la base de datos de estímulos en {_ESTIMULOS_JSON}"
        )
    
    norm_nivel = canonical(nivel)
    norm_asig = canonical(asignatura)
    if asignatura in SUBJECT_MAPS:
        norm_asig = canonical(SUBJECT_MAPS[asignatura].get("banco", asignatura))

    try:
        data = json.loads(_ESTIMULOS_JSON.read_text("utf-8"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error leyendo base de datos JSON: {str(e)}")

    # Lista negra de metadatos pedagógicos para evitar alucinaciones institucionales
    blacklist = [
        "aprendizajes esperados", "objetivo de aprendizaje", "tabla de especificaciones",
        "clave de respuesta", "pauta de", "eje temático", "habilidad evaluada",
        "indicadores de evaluación", "estándares de aprendizaje", "bases curriculares",
        "niveles de aprendizaje", "simce", "agencia de calidad", "mineduc"
    ]

    filtrados = []
    for item in data:
        contenido_limpio = item.get("contenido_real", "").lower()
        contains_garbage = any(bad_word in contenido_limpio for bad_word in blacklist)
        
        db_asig = canonical(item.get("asignatura", ""))
        match_asig = (db_asig in norm_asig) or (norm_asig in db_asig)
        db_nivel = canonical(item.get("nivel", ""))
        match_nivel = (db_nivel == norm_nivel)
        
        if match_asig and match_nivel and item.get("imagen_adjunta") and item.get("tipo") != "imagen_adjunta_solo" and not contains_garbage:
            filtrados.append(item)
    
    if not filtrados:
        raise HTTPException(
            status_code=404, 
            detail=f"No se encontraron estímulos para '{asignatura}' (canónigo: {norm_asig}) y nivel '{nivel}' (canónigo: {norm_nivel}). Revisa el JSON oficial."
        )
        
    return random.choice(filtrados)

# --- Helpers Exportación ---

def _safe_filename(name: str) -> str:
    ascii_name = unicodedata.normalize("NFKD", name).encode("ASCII", "ignore").decode()
    return re.sub(r"[^\w\s.-]", "", ascii_name).replace(" ", "_") or "documento"

def _docx_headers(filename: str) -> dict:
    safe = _safe_filename(filename)
    utf8 = quote(filename.replace(" ", "_"))
    return {
        "Content-Disposition": f'attachment; filename="{safe}"; filename*=UTF-8\'\'{utf8}',
        "Access-Control-Expose-Headers": "Content-Disposition",
    }

def _add_profeic_header(doc: Document, asignatura: str, nivel: str, tipo: str) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.2)
    t.columns[1].width = Inches(5.3)
    cell_logo = t.cell(0, 0)
    logo_path = _ASSETS_DIR / "logo_profeic.svg.png"
    p_logo = cell_logo.paragraphs[0]
    run_logo = p_logo.add_run()
    if logo_path.exists():
        try: run_logo.add_picture(str(logo_path), width=Inches(1.0))
        except: run_logo.add_text("ProfeIC")
    else: run_logo.add_text("ProfeIC")
    
    cell_info = t.cell(0, 1)
    p_info = cell_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_info.add_run("PROFE IC\n"); r1.bold = True; r1.font.size = Pt(12)
    r2 = p_info.add_run("Instrumento generado con IA · ProfeIC\n"); r2.font.size = Pt(8)
    r3 = p_info.add_run(f"Asignatura: {asignatura} | Nivel: {nivel} | {tipo}"); r3.font.size = Pt(9)
    doc.add_paragraph()

# --- Generación Segmentada (Prompt Chaining) ---

async def generar_pregunta_individual(skill: str, numero: int, nivel: str, asignatura: str, indicadores_texto: str, contexto_referencia: str, es_lenguaje: bool, historial_preguntas: str, foco_actual: str) -> dict:
    prompt_base = f"""Eres un psicómetra experto del MINEDUC para la asignatura {asignatura} ({nivel}).
Tu tarea es generar UNA (1) única pregunta de selección múltiple (A,B,C,D) que evalúe la habilidad de '{skill}'.

OBLIGATORIO: Esta pregunta debe enfocarse EXCLUSIVAMENTE en esta dimensión de análisis: {foco_actual}.

Contexto de referencia:
{contexto_referencia}

REGLAS ESTRICTAS:
1. La pregunta debe evaluar EXCLUSIVAMENTE la habilidad cognitiva: '{skill}' a través del prisma de: '{foco_actual}'.
"""
    if historial_preguntas:
        prompt_base += f"\nATENCIÓN - PREGUNTAS YA GENERADAS PARA ESTE TEXTO:\n{historial_preguntas}\nDEBES generar una pregunta completamente distinta evaluando '{skill}' con foco en '{foco_actual}'.\n"

    if not es_lenguaje:
        prompt_base += f"2. La pregunta DEBE evaluar conocimiento disciplinario real basado en los indicadores: '{indicadores_texto}'.\n"
    else:
        prompt_base += "2. La pregunta debe evaluar habilidades de comprensión lectora basadas en el texto principal.\n"
        
    prompt_base += f"""3. Diseña 1 alternativa correcta y 3 distractores.
4. REGLA DE ORO: Tienes estrictamente PROHIBIDO inventar o agregar textos de lectura adicionales. Debes generar la pregunta basándote ÚNICA y EXCLUSIVAMENTE en el contexto de referencia proporcionado arriba.
5. Retorna UNICAMENTE un objeto JSON válido con la siguiente estructura exacta:
{{
  "numero": {numero},
  "habilidad_medida": "{skill}",
  "enunciado": "Texto de la pregunta",
  "alternativas": {{
    "A": "Opción A",
    "B": "Opción B",
    "C": "Opción C",
    "D": "Opción D"
  }},
  "correcta": "Letra de la correcta",
  "justificacion": "Explicación detallada..."
}}
"""
    try:
        res = await ai_client.aio.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt_base,
            config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.3)
        )
        data = extract_json(res.text)
        if isinstance(data, dict):
            p = data
            p["numero"] = numero
            if "habilidad_medida" not in p and "habilidad" in p: p["habilidad_medida"] = p["habilidad"]
            p["habilidad_medida"] = skill
            if "pregunta" in p and "enunciado" not in p: p["enunciado"] = p["pregunta"]
            if "enunciado" not in p: p["enunciado"] = "Sin enunciado"
            if "justificacion" not in p: p["justificacion"] = "Sin justificación."
            if "correcta" not in p: p["correcta"] = "A"
            if "alternativas" not in p: p["alternativas"] = {"A": "Opción A", "B": "Opción B", "C": "Opción C", "D": "Opción D"}
            
            alts = p.get("alternativas")
            if isinstance(alts, list):
                new_alts = {}
                for i, val in enumerate(alts[:4]):
                    key = chr(65 + i)
                    new_alts[key] = re.sub(r"^[a-d][\)\.]\s*", "", str(val), flags=re.IGNORECASE)
                p["alternativas"] = new_alts
            elif isinstance(alts, dict):
                new_alts = {}
                for k, v in alts.items():
                    norm_k = str(k).upper().strip()
                    if norm_k in ["A", "B", "C", "D"]:
                        new_alts[norm_k] = re.sub(r"^[a-d][\)\.]\s*", "", str(v), flags=re.IGNORECASE)
                if new_alts: p["alternativas"] = new_alts
            return p
        else:
            return _fallback_question(numero, skill)
    except Exception as e:
        print(f"Error generating individual question: {e}")
        return _fallback_question(numero, skill)

def _fallback_question(numero: int, skill: str) -> dict:
    return {
        "numero": numero,
        "habilidad_medida": skill,
        "enunciado": "Hubo un error al generar esta pregunta.",
        "alternativas": {"A": "Error A", "B": "Error B", "C": "Error C", "D": "Error D"},
        "correcta": "A",
        "justificacion": "Error de generación de la IA o Timeout."
    }

# --- Generación de Bloque ---

async def generar_un_bloque(index: int, block_skills: List[str], nivel: str, asignatura: str, indicadores_texto: str, posibles_estimulos: List[Dict[str, Any]], pregunta_start_num: int) -> BloqueEnsayo:
    # --- Verificación estricta de indicadores ---
    if not indicadores_texto or indicadores_texto.strip() == "Marco curricular nacional general.":
        raise ValueError(f"Abordando: No se encontraron indicadores curriculares para {asignatura} en {nivel}. "
                         f"No podemos generar preguntas disciplinarias de calidad sin esta información.")

    estimulo_seleccionado = None
    if posibles_estimulos:
        idx_est = index % len(posibles_estimulos)
        estimulo_seleccionado = posibles_estimulos[idx_est]

    es_lenguaje = "lengua" in asignatura.lower() or "lenguaje" in asignatura.lower()

    if estimulo_seleccionado:
        # --- NUEVO: BYPASS PARA TEXTOS REALES ---
        if "contenido_real" in estimulo_seleccionado and estimulo_seleccionado["contenido_real"]:
            # Si hay contenido_real, usamos el texto directo (Bypass Fase 1)
            titulo = estimulo_seleccionado.get("titulo_interno", "Texto")
            autor = estimulo_seleccionado.get("autor_real", "Anónimo")
            contenido = estimulo_seleccionado["contenido_real"]
            
            texto_base = f"{titulo} - {autor}\n\n{contenido}"
            tipo_base = estimulo_seleccionado.get("tipo", "texto_literario_real")
            ruta_img = None
            contexto_referencia = f"TEXTO REAL PARA GENERAR PREGUNTAS: '{texto_base}'"
        else:
            # Lógica existente para estímulos basados en imágenes
            tipo_base = estimulo_seleccionado.get("tipo", "imagen_educativa")
            texto_base = "Observa la siguiente imagen y responde las preguntas:"
            descripcion_ia = estimulo_seleccionado.get("descripcion_oculta_ia", "")
            ruta_img = f"imgs/{estimulo_seleccionado.get('nombre_archivo_imagen')}"
            contexto_referencia = f"IMAGEN (No generes preguntas sobre la imagen literal, úsala como contexto disciplinario): '{descripcion_ia}'"
    else:
        # Fase 1: Generación de Estímulo con Gemini (Ciencias, Historia o sin banco)
        schema_instruccion = 'REGLA: Debes retornar UNICAMENTE un JSON con dos llaves: "tipo" y "contenido". El "contenido" debe tener el texto del estímulo.'
        if es_lenguaje:
            prompt_estimulo = f"Genera un estímulo textual inédito (narrativo o informativo). Nivel: {nivel}. Asignatura: {asignatura}. Basado en: {indicadores_texto}. 300 palabras. {schema_instruccion}"
        else:
            prompt_estimulo = f"Genera un breve estímulo disciplinario inédito (ej. un caso histórico, problema matemático, o experimento científico). Nivel: {nivel}. Asignatura: {asignatura}. Basado estrictamente en: {indicadores_texto}. Máximo 150 palabras. {schema_instruccion}"
            
        res_estimulo = await ai_client.aio.models.generate_content(model='gemini-2.5-flash', contents=prompt_estimulo, config=types.GenerateContentConfig(response_mime_type="application/json"))
        estimulo_data = extract_json(res_estimulo.text)
        if isinstance(estimulo_data, dict):
            texto_base = estimulo_data.get("contenido", "")
            tipo_base = estimulo_data.get("tipo", "texto_narrativo" if es_lenguaje else "caso_disciplinario")
        else:
            texto_base = str(estimulo_data)
            tipo_base = "texto_narrativo" if es_lenguaje else "caso_disciplinario"
        ruta_img = None
        contexto_referencia = f"ESTÍMULO DISCIPLINARIO GENERADO: '{texto_base}'"

    # Fase 2: Generación CONCURRENTE de Preguntas Individuales con Sincronización Taxonómica
    
    # Definición de Matrices de Criterios (Sincronización Habilidad-Foco)
    if es_lenguaje:
        matriz_criterios = [
            {"habilidad": "Interpretar y relacionar", "foco": "Propósito comunicativo o intención central del autor"},
            {"habilidad": "Interpretar y relacionar", "foco": "Análisis del estado psicológico, motivaciones o tono de los personajes/hablante"},
            {"habilidad": "Interpretar y relacionar", "foco": "Significado inferencial de vocabulario o expresiones en contexto"},
            {"habilidad": "Reflexionar", "foco": "Análisis de figuras literarias, simbolismo o recursos retóricos"},
            {"habilidad": "Reflexionar", "foco": "Relación estructural entre distintas partes del texto y su efecto"},
            {"habilidad": "Localizar", "foco": "Identificación de información explícita, datos precisos o secuencias de acciones literales"}
        ]
    elif "historia" in asignatura.lower():
        matriz_criterios = [
            {"habilidad": "Pensamiento temporal", "foco": "Contexto histórico o geográfico específico del suceso"},
            {"habilidad": "Pensamiento crítico", "foco": "Relación causa-efecto entre eventos o procesos sociales"},
            {"habilidad": "Análisis de fuentes", "foco": "Perspectivas de los actores históricos involucrados"},
            {"habilidad": "Pensamiento espacial", "foco": "Impacto del entorno geográfico en el desarrollo histórico"},
            {"habilidad": "Análisis de fuentes", "foco": "Interpretación de evidencia material o documental"}
        ]
    else: # Ciencias, Matemática, etc.
        matriz_criterios = [
            {"habilidad": "Razonamiento", "foco": "Análisis de variables o evidencia científica del caso"},
            {"habilidad": "Aplicación", "foco": "Interpretación de datos, tablas o gráficos presentados"},
            {"habilidad": "Conocimiento", "foco": "Relación entre conceptos teóricos y aplicaciones prácticas"},
            {"habilidad": "Razonamiento", "foco": "Evaluación de hipótesis o validez de métodos experimentales"},
            {"habilidad": "Evaluación", "foco": "Implicancias éticas, sociales o ambientales del fenómeno"}
        ]

    tareas = []
    for i in range(len(block_skills)):
        criterio = matriz_criterios[i % len(matriz_criterios)]
        habilidad_actual = criterio["habilidad"]
        foco_actual = criterio["foco"]
        
        # Creamos la tarea (corrutina) vinculando Habilidad y Foco
        tarea = generar_pregunta_individual(
            habilidad_actual, 
            pregunta_start_num + i, 
            nivel, 
            asignatura, 
            indicadores_texto, 
            contexto_referencia, 
            es_lenguaje, 
            "", 
            foco_actual
        )
        tareas.append(tarea)
    
    # Ejecución simultánea de todas las preguntas del bloque
    # return_exceptions=True asegura que si una falla, las demás continúen
    resultados = await asyncio.gather(*tareas, return_exceptions=False)
    
    lista_preguntas = []
    for res in resultados:
        if isinstance(res, dict):
            lista_preguntas.append(res)
        else:
            # En caso de error inesperado en una corrutina
            print(f"Error crítico en generación concurrente: {res}")
            lista_preguntas.append(_fallback_question(pregunta_start_num, "Error de Sistema"))
    
    return BloqueEnsayo(tipo_estimulo=tipo_base, contenido_estimulo=texto_base, preguntas=lista_preguntas, ruta_local=ruta_img)

# --- Endpoints ---

@router.post("/generate", response_model=SimceResponse)
async def generar_simce(req: SimceRequest, current_user_id: Optional[str] = Depends(get_current_user_id_optional)):
    try:
        # MODO MINI-ENSAYO DE PRUEBA: Forzamos 5 preguntas para evitar alucinaciones
        req.cantidad_preguntas = 5
        
        try:
            distribucion_total = calcular_distribucion(req.asignatura, req.nivel, req.cantidad_preguntas)
        except ValueError:
            habilidades_fallback = ["Localizar", "Interpretar y relacionar", "Reflexionar"]
            base_count = req.cantidad_preguntas // len(habilidades_fallback)
            distribucion_total = {h: base_count for h in habilidades_fallback}
            for i in range(req.cantidad_preguntas % 3): distribucion_total[habilidades_fallback[i]] += 1
        
        skills_pool = []
        for hab, cant in distribucion_total.items(): skills_pool.extend([hab] * cant)

        indicadores_texto = ""
        if SUPABASE_URL and SUPABASE_KEY:
            try:
                supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
                res = supabase.table("curriculum_indicadores").select("indicador").eq("nivel", req.nivel).eq("asignatura", req.asignatura).execute()
                if res.data: indicadores_texto = "\n- ".join([item["indicador"] for item in res.data[:20]])
            except: pass
        if not indicadores_texto: indicadores_texto = get_indicadores_locales(req.asignatura, req.nivel)
        if not indicadores_texto: indicadores_texto = "Marco curricular nacional general."

        tareas = []
        for i in range((req.cantidad_preguntas + 4) // 5):
            start = i * 5
            block_skills = skills_pool[start:start+5]
            if not block_skills: break
            
            # Obtener un estímulo específico para este bloque
            estimulo_bloque = _get_random_stimulus(req.asignatura, req.nivel)
            
            async def generar_con_estimulo(idx, skills, n, a, ind_txt, est_obj, start_num):
                bloque = await generar_un_bloque(idx, skills, n, a, ind_txt, [est_obj], start_num)
                # Inyectar estímulo en la primera pregunta del bloque
                if bloque.preguntas:
                    bloque.preguntas[0].estimulo_texto = est_obj.get("contenido_real") or bloque.contenido_estimulo
                    bloque.preguntas[0].estimulo_imagen = est_obj.get("imagen_adjunta") or bloque.ruta_local
                return bloque

            tareas.append(generar_con_estimulo(i, block_skills, req.nivel, req.asignatura, indicadores_texto, estimulo_bloque, 1 + start))

        ensayo_final = await asyncio.gather(*tareas)
        todas_las_preguntas = []
        for b in ensayo_final: todas_las_preguntas.extend(b.preguntas)
        
        # Extraer primer estímulo como global (Compatibilidad Frontend)
        est_global_txt = ensayo_final[0].preguntas[0].estimulo_texto if (ensayo_final and ensayo_final[0].preguntas) else None
        est_global_img = ensayo_final[0].preguntas[0].estimulo_imagen if (ensayo_final and ensayo_final[0].preguntas) else None

        return SimceResponse(
            titulo=f"Ensayo SIMCE — {req.asignatura} — {req.nivel}",
            asignatura=req.asignatura,
            nivel=req.nivel,
            cantidad_preguntas=len(todas_las_preguntas),
            modo=req.modo,
            blueprint_resumen=str(distribucion_total),
            estimulo_texto=est_global_txt,
            estimulo_imagen=est_global_img,
            ensayo=list(ensayo_final), 
            preguntas=todas_las_preguntas
        )
    except Exception as e:
        print(traceback.format_exc()); raise HTTPException(status_code=500, detail=str(e))

@router.post("/generate/download-cuadernillo")
async def download_cuadernillo(req: DescargaCuadernilloRequest, current_user_id: Optional[str] = Depends(get_current_user_id_optional)):
    try:
        # 0. Unificar fuente de preguntas (si viene como plana o como bloques)
        preguntas_finales = []
        if req.preguntas:
            preguntas_finales = req.preguntas
        elif req.ensayo:
            for b in req.ensayo:
                preguntas_finales.extend(b.preguntas)
        
        if not preguntas_finales:
            raise HTTPException(status_code=400, detail="No se encontraron preguntas para elevar al cuadernillo.")

        # 1. Descubrimiento robusto de la Raíz del Proyecto
        current_dir = Path(__file__).resolve()
        while current_dir.name != 'backend' and current_dir.parent != current_dir:
            current_dir = current_dir.parent
        project_root = current_dir.parent # La carpeta principal PROFEIC_...

        doc = Document()
        tipo_doc = "Pauta Docente" if req.incluir_clave else "Cuadernillo del Alumno"
        _add_profeic_header(doc, req.asignatura, req.nivel, tipo_doc)

        h = doc.add_heading(req.titulo, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs: run.font.color.rgb = RGBColor(27, 60, 115)

        p_sub = doc.add_paragraph(f"{'Simulación SIMCE' if req.modo == 'simce' else 'Ensayo Formativo'} · {req.nivel}")
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.runs[0].italic = True

        # Nueva Estética SIMCE: Fuente base
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # 1.5 Estímulo GLOBAL (si se envía en el request)
        if req.estimulo_imagen:
            imagen_path = project_root / "frontend" / "public" / "imgs_estimulos" / req.estimulo_imagen
            try:
                if imagen_path.exists():
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph().add_run().add_picture(str(imagen_path), width=Inches(5.0))
                else:
                    print(f"ERROR: Imagen global no encontrada en {imagen_path}")
            except Exception as e:
                print(f"Error imagen global: {e}")

        if req.estimulo_texto:
            t_est = doc.add_table(rows=1, cols=1)
            t_est.style = "Table Grid"
            cell = t_est.cell(0, 0)
            cell.text = req.estimulo_texto
            shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            doc.add_paragraph()

        # 2. Impresión de Preguntas con Estímulos Intercalados
        for p in preguntas_finales:
            # A. DIAGRAMACIÓN MULTIMODAL (IMAGEN)
            if p.estimulo_imagen:
                # Búsqueda robusta basada en project_root
                nombre_limpio = p.estimulo_imagen.replace("imgs/", "").replace("imgs_estimulos/", "")
                ruta_img = project_root / "frontend" / "public" / "imgs_estimulos" / nombre_limpio
                
                if ruta_img.exists():
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(str(ruta_img), width=Inches(5.0))
                    doc.add_paragraph()
                else:
                    print(f"ERROR: Imagen pregunta no encontrada en {ruta_img}")

            # B. DIAGRAMACIÓN MULTIMODAL (RECUADRO GRIS TEXTO)
            if p.estimulo_texto:
                t_est = doc.add_table(rows=1, cols=1)
                t_est.style = "Table Grid"
                cell = t_est.cell(0, 0)
                cell.text = p.estimulo_texto
                shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                doc.add_paragraph()

            # C. IMPRESIÓN DE LA PREGUNTA
            p_enun = doc.add_paragraph()
            run_num = p_enun.add_run(f"{p.numero}. "); run_num.bold = True
            p_enun.add_run(p.enunciado)
            
            # Habilidad badge (sutil)
            p_hab = doc.add_paragraph()
            run_hab = p_hab.add_run(f"   [{p.habilidad_medida}]")
            run_hab.italic = True; run_hab.font.size = Pt(8); run_hab.font.color.rgb = RGBColor(120, 120, 120)

            # Alternativas con sangría
            for letra, alt_info in p.alternativas.items():
                # Extraer texto si alt_info es un dict {"texto": "...", ...}
                texto_final = alt_info.get("texto", str(alt_info)) if isinstance(alt_info, dict) else str(alt_info)
                p_alt = doc.add_paragraph(f"{letra}) {texto_final}")
                p_alt.paragraph_format.left_indent = Inches(0.5)
                if req.incluir_clave and letra == p.correcta:
                    p_alt.runs[0].bold = True
                    p_alt.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            doc.add_paragraph()

        if req.incluir_clave:
            doc.add_page_break()
            doc.add_heading("Pauta de Respuestas", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            for p in preguntas_finales:
                doc.add_paragraph(f"Pregunta {p.numero}: {p.correcta}").runs[0].bold = True
                doc.add_paragraph(p.justificacion)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=_docx_headers(f"Cuadernillo_{req.asignatura}.docx"))
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))

import uuid
from services.omr_template_service import OMRTemplateGenerator

@router.post("/generate/download-omr")
async def download_omr(req: DescargaOMRRequest, current_user_id: Optional[str] = Depends(get_current_user_id_optional)):
    try:
        # Generar un ID único para personalizar el código QR de esta descarga
        eval_id = f"SIMCE-{uuid.uuid4().hex[:8].upper()}"
        
        generator = OMRTemplateGenerator()
        buf = io.BytesIO()
        
        logo_path = _ASSETS_DIR / "logo_profeic.svg.png"
        if not logo_path.exists():
            logo_path = None
        else:
            logo_path = str(logo_path)
            
        generator.generate_pdf(buf, eval_id, num_questions=req.cantidad_preguntas, logo_path=logo_path)
        buf.seek(0)
        
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="Hoja_OMR_{req.asignatura}_{eval_id}.pdf"',
                "Access-Control-Expose-Headers": "Content-Disposition",
            }
        )
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))
